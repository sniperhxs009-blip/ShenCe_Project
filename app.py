# ========== 调试定位：运行 streamlit run app.py 后，看终端打印停在哪一步 ==========
import sys
def _dbg(msg):
    print(f">>> [{msg}]", flush=True)
_dbg("1-开始导入")

import streamlit as st
_dbg("2-streamlit")
from openai import OpenAI
_dbg("3-openai")
import json
import os
_dbg("4-json,os")
import plotly.graph_objects as go
_dbg("5-plotly")
import requests
import random
from datetime import datetime
import pandas as pd
_dbg("6-pandas")
from fpdf import FPDF
import io
import time
import re
import math

# 引入外部文件解析模块（与 file-report-generator 保持一致的解析逻辑）
EXTERNAL_PARSER_AVAILABLE = False
try:
    FILE_PARSER_ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "file-report-generator")
    if os.path.isdir(FILE_PARSER_ROOT):
        if FILE_PARSER_ROOT not in sys.path:
            sys.path.append(FILE_PARSER_ROOT)
        try:
            from document_parser import parse_document as ext_parse_document
            EXTERNAL_PARSER_AVAILABLE = True
        except Exception:
            ext_parse_document = None  # type: ignore
    else:
        ext_parse_document = None  # type: ignore
except Exception:
    ext_parse_document = None  # type: ignore

_dbg("7-基础库完成")

# PDF 解析：优先 PyMuPDF（质量更好），其次 pypdf
PDF_AVAILABLE = False
PYMUPDF_AVAILABLE = False
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
    PDF_AVAILABLE = True
except ImportError:
    try:
        from pypdf import PdfReader
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

# Word 解析（可选依赖）
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# OCR 兜底（扫描版 PDF 文字极少时启用）
OCR_AVAILABLE = False
try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_AVAILABLE = True
except ImportError:
    pass
_dbg("8-PDF/DOCX可选依赖")

# --- 页面配置 ---
st.set_page_config(page_title="事件推演仿真系统", layout="wide", initial_sidebar_state="expanded")
_dbg("9-set_page_config")

# 恢复接近 Streamlit 默认的浅色背景，只保留少量卡片样式
st.markdown("""
<style>
.block-container { max-width: 98% !important; padding: 1rem 2% !important; }
.metric-card { 
    background-color: #f8f9fa; padding: 20px; border-radius: 10px; 
    border: 1px solid #e0e0e0; text-align: center;
}
.logic-box { 
    background-color: #f8f9fa; padding: 25px; border-left: 10px solid #58a6ff; 
    border-radius: 6px; margin: 20px 0; font-family: 'Consolas', monospace; 
    color: #555555; border: 1px solid #e0e0e0;
}
.report-card { 
    background-color: #ffffff; padding: 40px; border-radius: 15px; color: #1a1a1a !important;
    border-top: 15px solid #1f6feb; box-shadow: 0 10px 40px rgba(0,0,0,0.05); 
}
.agent-card {
    background-color: #ffffff; padding:15px; border-radius:10px; min-height:260px;
    border:1px solid #e0e0e0; margin-bottom:10px;
}
.resource-card {
    background-color: #ffffff; padding:15px; border-radius:10px;
    border:1px solid #e0e0e0; margin-bottom:10px;
}
</style>
""", unsafe_allow_html=True)
_dbg("10-CSS样式")

# --- 侧边栏 API ---
st.sidebar.header("🔑 API 配置")
openai_key = st.sidebar.text_input("DeepSeek API Key（兼容 OpenAI SDK）", type="password")
base_url = st.sidebar.text_input("Base URL", value="https://api.deepseek.com/v1")
serper_key = st.sidebar.text_input("Serper API Key", type="password")

client = None
if openai_key:
    client = OpenAI(api_key=openai_key, base_url=base_url)
_dbg("11-侧栏API区")

# --- 仿真模式与可复现性 ---
st.sidebar.header("🧪 仿真模式")
sim_mode = st.sidebar.selectbox(
    "模式",
    ["高保真（机制内核 + 可审计）", "叙事（AI 主导）"],
    help="高保真模式：指标与资源由机制模型计算，AI 仅做文字解读并受审计约束；叙事模式：主要由模型生成并辅以随机扰动。"
)
seed = st.sidebar.number_input("随机种子（可复现）", min_value=0, max_value=10_000_000, value=2026, step=1)
force_crisis_mode = st.sidebar.checkbox(
    "强制危机推演模式",
    value=False,
    help="勾选后忽略智能判断，始终按政府/社会危机推演流程执行（即使输入是续写小说等）"
)

# --- 高保真参数校准面板 ---
hf_default = {
    "shock_anxiety_k": 1.0,
    "shock_gap_k": 1.0,
    "shock_risk_k": 1.0,
    "shock_admin_k": 1.0,
    "coupling_anxiety_k": 1.0,
    "coupling_risk_k": 1.0,
    "trigger_risk": 80,
    "trigger_gap": 75,
    "trigger_admin": 25,
    "min_infra_to_execute": 25,
    "step_hours": 12.0,
    "decay_scale_panic": 0.6,
    "decay_scale_supply_shock": 1.2,
    "decay_scale_power_outage": 1.0,
    "decay_scale_comms_outage": 0.8,
    "decay_scale_finance_risk": 1.1,
    "decay_scale_governance_stress": 0.9,
}
with st.sidebar.expander("🎛️ 高保真参数校准（进阶）", expanded=False):
    st.caption("这些参数只影响高保真模式，用于调参/回放对比。")
    shock_anxiety_k = st.slider("冲击→焦虑 系数", 0.0, 2.0, float(hf_default["shock_anxiety_k"]), 0.05)
    shock_gap_k = st.slider("冲击→缺口 系数", 0.0, 2.0, float(hf_default["shock_gap_k"]), 0.05)
    shock_risk_k = st.slider("冲击→动荡 系数", 0.0, 2.0, float(hf_default["shock_risk_k"]), 0.05)
    shock_admin_k = st.slider("冲击→行政 系数", 0.0, 2.0, float(hf_default["shock_admin_k"]), 0.05)
    coupling_anxiety_k = st.slider("耦合→焦虑 系数", 0.0, 2.0, float(hf_default["coupling_anxiety_k"]), 0.05)
    coupling_risk_k = st.slider("耦合→动荡 系数", 0.0, 2.0, float(hf_default["coupling_risk_k"]), 0.05)
    trigger_risk = st.slider("动荡阈值触发", 50, 95, int(hf_default["trigger_risk"]), 1)
    trigger_gap = st.slider("缺口阈值触发", 50, 95, int(hf_default["trigger_gap"]), 1)
    trigger_admin = st.slider("低行政阈值触发", 5, 50, int(hf_default["trigger_admin"]), 1)
    min_infra_to_execute = st.slider("执行最低基础设施阈值", 0, 60, int(hf_default["min_infra_to_execute"]), 1)
    step_hours = st.slider("每步代表小时数", 1.0, 48.0, float(hf_default["step_hours"]), 1.0, help="用于把证据持续时间换算为衰减曲线；相同证据下，步长越大衰减越快。")
    auto_steps_by_duration = st.checkbox("按总时长自动计算步数（高保真）", value=False, help="开启后，高保真模式将忽略上方“演化步数”滑块，改用总时长/每步小时数计算步数。")
    dur_cols = st.columns(2)
    total_duration_hours = dur_cols[0].number_input("总推演时长（小时）", min_value=0, max_value=24*60, value=0, step=12, help="0 表示不启用；启用后步数=ceil(总时长/每步小时数)。")
    total_duration_days = dur_cols[1].number_input("总推演时长（天）", min_value=0.0, max_value=60.0, value=0.0, step=0.5, help="可选：用天输入总时长；若>0 将优先于“小时”。")
    st.caption("起止时间（可选，用于昼夜节律与审计对齐）")
    time_cols = st.columns(2)
    start_time_str = time_cols[0].text_input("开始时间（YYYY-MM-DD HH:MM）", value="", placeholder="如：2026-03-17 08:00")
    end_time_str = time_cols[1].text_input("结束时间（YYYY-MM-DD HH:MM）", value="", placeholder="如：2026-03-19 20:00")

    circadian = st.checkbox("启用昼夜节律（高保真）", value=True, help="启用后，每一步会映射到具体时刻，夜间焦虑/动荡更敏感，白天执行更强。")
    night_anxiety_boost = st.slider("夜间焦虑增益", 0.0, 1.0, 0.25, 0.05)
    night_risk_boost = st.slider("夜间动荡增益", 0.0, 1.0, 0.20, 0.05)
    night_exec_penalty = st.slider("夜间执行折扣", 0.0, 1.0, 0.15, 0.05)
    st.caption("因子衰减倍率（>1 衰减更慢，<1 衰减更快）")
    decay_scale_panic = st.slider("恐慌/舆情 衰减倍率", 0.2, 2.5, float(hf_default["decay_scale_panic"]), 0.05)
    decay_scale_supply_shock = st.slider("供应链冲击 衰减倍率", 0.2, 2.5, float(hf_default["decay_scale_supply_shock"]), 0.05)
    decay_scale_power_outage = st.slider("电力故障 衰减倍率", 0.2, 2.5, float(hf_default["decay_scale_power_outage"]), 0.05)
    decay_scale_comms_outage = st.slider("通信故障 衰减倍率", 0.2, 2.5, float(hf_default["decay_scale_comms_outage"]), 0.05)
    decay_scale_finance_risk = st.slider("金融风险 衰减倍率", 0.2, 2.5, float(hf_default["decay_scale_finance_risk"]), 0.05)
    decay_scale_governance_stress = st.slider("治理压力 衰减倍率", 0.2, 2.5, float(hf_default["decay_scale_governance_stress"]), 0.05)

hf_params = {
    "shock_anxiety_k": shock_anxiety_k,
    "shock_gap_k": shock_gap_k,
    "shock_risk_k": shock_risk_k,
    "shock_admin_k": shock_admin_k,
    "coupling_anxiety_k": coupling_anxiety_k,
    "coupling_risk_k": coupling_risk_k,
    "trigger_risk": trigger_risk,
    "trigger_gap": trigger_gap,
    "trigger_admin": trigger_admin,
    "min_infra_to_execute": min_infra_to_execute,
    "step_hours": step_hours,
    "auto_steps_by_duration": bool(auto_steps_by_duration),
    "total_duration_hours": int(total_duration_hours),
    "total_duration_days": float(total_duration_days),
    "start_time_str": start_time_str,
    "end_time_str": end_time_str,
    "circadian": bool(circadian),
    "night_anxiety_boost": float(night_anxiety_boost),
    "night_risk_boost": float(night_risk_boost),
    "night_exec_penalty": float(night_exec_penalty),
    "decay_scales": {
        "panic": decay_scale_panic,
        "supply_shock": decay_scale_supply_shock,
        "power_outage": decay_scale_power_outage,
        "comms_outage": decay_scale_comms_outage,
        "finance_risk": decay_scale_finance_risk,
        "governance_stress": decay_scale_governance_stress,
    },
}
_dbg("12-高保真参数")

# --- 会话状态 ---
init_keys = [
    "event", "facts", "matrix_history", "timeline", "report",
    "agents", "resources", "swan", "causal_chain", "playback_index",
    "reset_requested", "autoplay_running",
    "audit_log", "evidence_items", "infra_history", "hf_params",
    "evidence_factor_map",
    "decay_schedule", "time_axis",
    "uploaded_doc_text", "sim_phase", "stepping_current_step", "stepping_effective_steps",
    "stepping_matrix", "stepping_resources", "stepping_infra", "stepping_factors",
    "perspective_chat_history",
    "scenario_branches", "current_branch_view", "entity_graph", "history_scenarios",
]
for k in init_keys:
    if k not in st.session_state:
        if k in ["timeline", "matrix_history", "audit_log", "evidence_items", "infra_history", "perspective_chat_history", "scenario_branches", "history_scenarios"]:
            st.session_state[k] = []
        elif k in ["history_scenarios"]:
            st.session_state[k] = []
        elif k in ["reset_requested", "autoplay_running"]:
            st.session_state[k] = False
        elif k in ["stepping_current_step", "stepping_effective_steps"]:
            st.session_state[k] = 0
        elif k == "sim_phase":
            st.session_state[k] = ""
        else:
            st.session_state[k] = ""

st.session_state.hf_params = hf_params

if st.session_state.get("playback_index") is None or not isinstance(st.session_state.get("playback_index"), int):
    st.session_state.playback_index = 0
_dbg("13-会话状态初始化")

# --- 文档解析（种子上传）---
def _decode_text(raw: bytes, max_len: int = 60000) -> str:
    """TXT/MD 多编码尝试：utf-8、gbk、gb2312、big5，兼容中文文件。"""
    for enc in ("utf-8", "gbk", "gb2312", "big5", "utf-16", "utf-16-le"):
        try:
            return raw.decode(enc, errors="strict")[:max_len]
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")[:max_len]

def parse_uploaded_document(uploaded_file) -> str:
    """解析上传的 PDF、TXT、DOCX、MD，返回纯文本。
    优先使用与 file-report-generator 一致的解析逻辑，确保读取结果完全一致；
    若外部解析模块不可用，再回退到本地实现（PyMuPDF/pypdf/docx/OCR）。
    """
    if not uploaded_file:
        return ""
    try:
        name = (uploaded_file.name or "").lower()
        raw = uploaded_file.read()
        max_chars = 60000  # 本地兜底解析时的最大长度限制（外部解析不截断，以保持一致）
        max_pages = 80  # 页数过多会导致首次解析过慢

        # 1) 若存在外部解析模块（来自 file-report-generator），则走与独立网页完全一致的解析路径
        if EXTERNAL_PARSER_AVAILABLE and name:
            ext = os.path.splitext(name)[1].lower()
            if ext in [".pdf", ".docx", ".doc", ".txt"]:
                tmp_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_uploaded_tmp")
                os.makedirs(tmp_dir, exist_ok=True)
                # 简单安全的文件名处理，避免路径注入
                base = os.path.basename(name)
                safe = re.sub(r"[^0-9a-zA-Z\u4e00-\u9fff\._-]", "_", base)
                tmp_path = os.path.join(tmp_dir, safe or "uploaded")
                try:
                    with open(tmp_path, "wb") as f:
                        f.write(raw)
                    # 与 file-report-generator 中的行为保持一致：不额外截断内容
                    return str(ext_parse_document(tmp_path) or "")
                finally:
                    try:
                        os.remove(tmp_path)
                    except Exception:
                        pass

        # 2) 外部解析不可用时，回退到原有本地解析逻辑
        if name.endswith(".txt"):
            return _decode_text(raw, max_chars)
        if name.endswith(".md") or name.endswith(".markdown"):
            return _decode_text(raw, max_chars)

        if name.endswith(".pdf") and PDF_AVAILABLE:
            text_parts = []
            if PYMUPDF_AVAILABLE:
                try:
                    doc = fitz.open(stream=raw, filetype="pdf")
                    n_pages = min(doc.page_count, max_pages)
                    for i in range(n_pages):
                        page = doc.load_page(i)
                        t = page.get_text("text", sort=True) or ""
                        if t.strip():
                            text_parts.append(t)
                    doc.close()
                except Exception:
                    pass
            if not text_parts and PDF_AVAILABLE:
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(io.BytesIO(raw))
                    for i, p in enumerate(reader.pages):
                        if i >= max_pages:
                            break
                        t = p.extract_text() or ""
                        if t.strip():
                            text_parts.append(t)
                except Exception:
                    pass
            out = "\n\n".join(text_parts)
            # OCR 兜底：文字极少（如扫描版）时尝试 OCR
            if OCR_AVAILABLE and len(out.strip()) < 500:
                try:
                    images = convert_from_bytes(raw, first_page=1, last_page=min(5, max_pages), dpi=120)
                    ocr_parts = []
                    for img in images:
                        ocr_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
                        if ocr_text.strip():
                            ocr_parts.append(ocr_text)
                    if ocr_parts:
                        out = "\n\n".join(ocr_parts)[:max_chars]
                except Exception:
                    pass
            return out[:max_chars] if out else ""
        if name.endswith(".pdf") and not PDF_AVAILABLE:
            return "[PDF 解析需要安装 pypdf 或 pymupdf：pip install pymupdf]"

        if name.endswith(".docx") and DOCX_AVAILABLE:
            doc = DocxDocument(io.BytesIO(raw))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for tbl in doc.tables:
                for row in tbl.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        parts.append(" | ".join(row_text))
                parts.append("")  # 表格行间空行
            return "\n".join(parts)[:max_chars]
        if (name.endswith(".docx") or name.endswith(".doc")) and not DOCX_AVAILABLE:
            return "[Word 解析需要安装 python-docx：pip install python-docx]"
        if name.endswith(".doc"):
            return "[仅支持 .docx 格式，请另存为 docx 后上传]"
    except Exception as e:
        return f"[文档解析异常] {str(e)[:120]}"
    return ""

def analyze_uploaded_text(text: str) -> dict:
    """
    对上传文档的全文内容进行结构化分析，用于后续研判。
    逻辑参考 file-report-generator 中的 analyze_content。
    """
    if not text or not str(text).strip():
        return {"error": "文档内容为空"}

    text = str(text).strip()
    lines = [l for l in text.split("\n") if l.strip()]

    char_count = len(text)
    word_count = len(re.findall(r"[\u4e00-\u9fff\w]+", text))
    line_count = len(lines)
    paragraph_count = len(re.split(r"\n\s*\n", text))

    headings = []
    for line in lines[:50]:
        line = line.strip()
        if 2 <= len(line) <= 50 and not line.endswith(("。", "，")):
            headings.append(line)

    words = re.findall(r"[\u4e00-\u9fff]{2,}|[a-zA-Z]{3,}", text)
    word_freq: dict[str, int] = {}
    for w in words:
        word_freq[w] = word_freq.get(w, 0) + 1
    top_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:15]

    summary = text[:300].replace("\n", " ")
    if len(text) > 300:
        summary += "..."

    content_preview = text[:500]

    return {
        "char_count": char_count,
        "word_count": word_count,
        "line_count": line_count,
        "paragraph_count": paragraph_count,
        "headings": headings[:10],
        "keywords": [{"word": w, "count": c} for w, c in top_keywords],
        "summary": summary,
        "content_preview": content_preview,
    }

def generate_uploaded_doc_report(analysis: dict, filename: str, file_type: str) -> str:
    """
    基于全文分析结果生成“内容研判 + 建议”的综合报告。
    逻辑参考 file-report-generator 中的 generate_report（第二段实现）。
    """
    if not analysis or analysis.get("error"):
        return analysis.get("error", "文档内容为空或解析失败")

    keywords_top = [k["word"] for k in analysis.get("keywords", [])[:5]]
    theme_str = "、".join(keywords_top) if keywords_top else "（未识别到明显主题词）"
    para_count = int(analysis.get("paragraph_count", 0) or 0)
    word_count = int(analysis.get("word_count", 0) or 0)
    content_preview = analysis.get("content_preview", analysis.get("summary", "")) or ""

    raw_sentences = re.split(r"[。！？!?\\n]", content_preview)
    sentences = [s.strip() for s in raw_sentences if s.strip()]
    core_points = sentences[:6]

    risk_words = ["风险", "问题", "不足", "隐患", "挑战", "困难", "矛盾", "压力"]
    risk_sentences = [s for s in sentences if any(w in s for w in risk_words)][:5]

    report = f"""
# 文档内容分析研判报告（基于全文）

**生成时间**: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}

**文件名**: {filename or "未命名文档"}

**文件类型**: {file_type.upper() if file_type else "未知"}

---

## 一、内容概览与主旨判断

{analysis.get('summary', '')}

- 初步判断：文档重点围绕 **{theme_str}** 等议题展开。

---

## 二、基础规模与结构特征

> 下列数据用于理解文档篇幅与结构特征。

- 约 **{analysis.get('word_count', 0):,}** 个词汇，**{analysis.get('paragraph_count', 0)}** 个段落；
- 行数约 **{analysis.get('line_count', 0):,}** 行，文本总体篇幅属{"较短" if word_count < 200 else "中等" if word_count < 1500 else "偏长"}范围；
- 结构上{"层次较为清晰，适合进一步拆解重点内容" if para_count > 3 else "结构相对简洁，更适合快速阅读和提炼关键信息"}。

---

## 三、结构要点与章节线索

### 可能的章节或小节标题（按出现顺序）
"""
    for h in analysis.get("headings", [])[:8]:
        report += f"- {h}\n"

    report += """

---

## 四、关键词分布与关注焦点

| 关键词 | 出现次数 |
|--------|----------|
"""
    for kw in analysis.get("keywords", [])[:10]:
        report += f"| {kw['word']} | {kw['count']} |\n"

    report += """

从上述高频词可以看出，文档在以下方向上着墨较多：""" + (theme_str if theme_str else "当前文本中未能明显识别出高频主题词") + "。"

    report += """

---

## 五、核心观点与关键信息提炼

结合文档内容，梳理若干有代表性的要点，供快速理解与汇报使用：
"""
    if core_points:
        for idx, s in enumerate(core_points, 1):
            report += f"{idx}. {s}\n"
    else:
        report += "- 当前文档内容较为零散，暂无法自动提炼清晰的关键要点，请结合原文进行人工补充。\n"

    report += """

---

## 六、潜在风险、问题与影响研判

结合文中表述，对可能存在的风险点和问题进行初步识别（仅供参考）：
"""
    if risk_sentences:
        for idx, s in enumerate(risk_sentences, 1):
            report += f"{idx}. {s}\n"
        report += "\n> 建议对上述内容逐条梳理，纳入风险台账或整改清单。\n"
    else:
        report += (
            "- 在可识别的文本片段中，未出现明显的“风险 / 问题 / 不足”等关键词表述；\n"
            "- 仍建议结合业务实际，从目标达成度、资源保障、进度安排、外部环境等维度，进一步排查潜在风险。\n"
        )

    report += """

---

## 七、意见建议与行动方向

在不改变原文含义的前提下，根据文本主题与表述方式，提出若干可落地的参考建议：
"""
    if word_count < 300:
        report += (
            "1. 建议在现有基础上，进一步补充背景、目标和评估标准，使文档从“情况说明”升级为可直接指导工作的“任务清单”。\n"
            "2. 对文中涉及的关键工作事项，建议形成简要台账（事项—责任人—时间节点—预期结果），便于跟踪落实。\n"
        )
    elif word_count < 2000:
        report += (
            "1. 建议从文中高频主题（如：" + theme_str + "）出发，梳理 3–5 条核心目标，并与现有措施逐一对应。\n"
            "2. 对关键环节设置可量化的评价指标（时间、质量、风险、资源等），将“原则性表述”转化为“可检查事项”。\n"
            "3. 对可能存在的风险点，建议在文末单列“风险及应对措施”部分，明确预案和兜底安排。\n"
        )
    else:
        report += (
            "1. 文档内容较为详尽，建议在此基础上提炼 1–2 页的高层概要（包含背景、目标、路径、风险、结论），用于领导汇报或部门沟通。\n"
            "2. 将文中涉及的重点任务转化为项目化清单（里程碑—关键交付物—责任单位—资源需求），以便纳入整体工作计划管理。\n"
            "3. 对涉及多部门协同的事项，建议补充“协同机制”和“信息共享机制”的设计，避免执行过程出现断点或重复劳动。\n"
        )

    report += """

---

## 八、综合结论

总体来看，本文件在阐述“""" + theme_str + """”等方面具有一定的完整性和参考价值。
建议将上述“核心观点”“风险问题”“意见建议”作为二次加工的基础材料，用于形成更具针对性的决策支持报告或专题研判材料。
"""

    return report.strip()

def fetch_url_as_seed(url: str) -> str:
    """抓取网页内容作为种子，返回纯文本。"""
    if not url or not url.strip():
        return ""
    url = url.strip()
    if not url.startswith("http://") and not url.startswith("https://"):
        url = "https://" + url
    try:
        r = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0 (compatible; SHENCE/1.0)"})
        r.raise_for_status()
        html = r.text
        text = re.sub(r"<script[^>]*>[\s\S]*?</script>", "", html, flags=re.I)
        text = re.sub(r"<style[^>]*>[\s\S]*?</style>", "", text, flags=re.I)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text[:12000]
    except Exception as e:
        return f"[URL 抓取异常] {str(e)[:80]}"

# --- 核心工具函数 ---
def safe_json(res):
    try:
        return json.loads(res)
    except:
        fixed = res[res.find("{"):res.rfind("}")+1]
        try:
            return json.loads(fixed)
        except:
            return {
                "official": "数据异常", "citizen": "数据异常",
                "media": "数据异常", "audit": "数据异常"
            }

def clamp01(x: float) -> float:
    return max(0.0, min(1.0, float(x)))

def clamp100(x: float) -> int:
    return int(max(0, min(100, round(float(x)))))

def _url_host(url: str) -> str:
    try:
        # extremely small parser; avoids extra deps
        u = (url or "").strip().lower()
        u = u.replace("https://", "").replace("http://", "")
        return u.split("/")[0]
    except Exception:
        return ""

def score_evidence_strength(item: dict) -> dict:
    """
    证据强度评分（0..1），用于因子权重（减少“随便命中关键词就很强”的问题）。
    这是启发式规则：可复现、可审计、可逐步替换为更严谨的打分器。
    """
    title = (item.get("title") or "").strip()
    snippet = (item.get("snippet") or "").strip()
    link = (item.get("link") or "").strip()
    host = _url_host(link)
    text = f"{title} {snippet}"

    length_score = clamp01(len(snippet) / 220.0)  # 片段越长信息越多
    has_numbers = 1.0 if re.search(r"\d", text) else 0.0
    has_time = 1.0 if re.search(r"(20\d{2}|19\d{2}|48\s*小时|7\s*天|周|月)", text) else 0.0

    # 域名可信度（保守：只给少量加分，避免“域名决定真伪”）
    host_bonus = 0.0
    if any(x in host for x in ["gov", "edu", "who.int", "un.org", "worldbank.org", "oecd.org", "imf.org"]):
        host_bonus = 0.15
    elif any(x in host for x in ["wikipedia.org", "reuters.com", "bbc.", "nytimes.com"]):
        host_bonus = 0.10

    strength = clamp01(0.15 + 0.35 * length_score + 0.20 * has_numbers + 0.15 * has_time + host_bonus)
    return {
        "host": host,
        "strength": round(float(strength), 2),
        "signals": {
            "length_score": round(float(length_score), 2),
            "has_numbers": bool(has_numbers),
            "has_time": bool(has_time),
            "host_bonus": round(float(host_bonus), 2),
        },
    }

def extract_evidence_struct(item: dict) -> dict:
    """
    结构化抽取（确定性规则）：事件类型/时间尺度/主题标签/数值片段。
    """
    text = f"{item.get('title','')} {item.get('snippet','')}".lower()
    tags = []
    if any(k in text for k in ["停电", "断电", "blackout", "power outage"]):
        tags.append("电力故障")
    if any(k in text for k in ["断网", "通信", "internet outage", "telecom"]):
        tags.append("通信故障")
    if any(k in text for k in ["供应链", "物流", "运输", "shortage", "supply chain"]):
        tags.append("供应链冲击")
    if any(k in text for k in ["挤兑", "bank run", "liquidity", "兑付"]):
        tags.append("金融风险")
    if any(k in text for k in ["恐慌", "抢购", "hoard", "panic", "rumor", "谣言"]):
        tags.append("舆情/恐慌")
    if any(k in text for k in ["宵禁", "curfew", "封控", "lockdown"]):
        tags.append("强制管控")

    nums = re.findall(r"(?:20\d{2}|19\d{2}|\d+\s*(?:小时|天|周|%|万人|万|亿|千|百))", item.get("snippet","") or "")
    time_scale = "未知"
    if re.search(r"48\s*小时|\d+\s*小时", text):
        time_scale = "小时级"
    elif re.search(r"\d+\s*天|周", text):
        time_scale = "天-周级"
    elif re.search(r"月|季度|年|20\d{2}|19\d{2}", text):
        time_scale = "月-年级"

    # 持续时间粗抽取（用于冲击衰减）
    duration_hours = None
    m_h = re.search(r"(\d+)\s*小时", text)
    m_d = re.search(r"(\d+)\s*天", text)
    m_w = re.search(r"(\d+)\s*周", text)
    if m_h:
        duration_hours = int(m_h.group(1))
    elif m_d:
        duration_hours = int(m_d.group(1)) * 24
    elif m_w:
        duration_hours = int(m_w.group(1)) * 24 * 7

    return {
        "tags": tags,
        "time_scale": time_scale,
        "numbers": nums[:6],
        "duration_hours": duration_hours,
    }

def compute_decay_schedule(evidence_items: list[dict], steps: int, step_hours: float, decay_scales: dict) -> dict:
    """
    根据证据中的持续时间/时间尺度，为各冲击因子生成随步数衰减的 multiplier（0..1）。
    规则：half_life_steps 与持续时间正相关；若未知则给中等半衰期。
    """
    # 从证据里估计一个总体持续时间（小时）
    durations = [it.get("duration_hours") for it in (evidence_items or []) if isinstance(it.get("duration_hours"), int)]
    if durations:
        dur_h = sorted(durations)[len(durations) // 2]  # 中位数
    else:
        # 用 time_scale 兜底
        scales = [it.get("time_scale") for it in (evidence_items or []) if it.get("time_scale")]
        if "小时级" in scales:
            dur_h = 36
        elif "天-周级" in scales:
            dur_h = 24 * 5
        elif "月-年级" in scales:
            dur_h = 24 * 30
        else:
            dur_h = 24 * 3

    # 将持续时间映射到半衰期（步）
    step_hours = float(step_hours or 12.0)
    dur_steps = max(1.0, dur_h / step_hours)
    half_life = max(1.0, dur_steps / 2.0)

    def decay(t: int) -> float:
        # t 从 1..steps；t 越大 multiplier 越小
        # multiplier = 0.5 ** ((t-1)/half_life)
        return float(0.5 ** ((max(0, t - 1)) / half_life))

    # 证据 tags 自动微调：若证据明确包含某类故障，对应因子半衰期更长（衰减更慢）
    tags = []
    for it in evidence_items or []:
        tags += (it.get("tags") or [])
    tags = list({t for t in tags if t})

    scales = decay_scales or {}
    def factor_half_life(f: str) -> float:
        s = float(scales.get(f, 1.0))
        boost = 1.0
        if f == "power_outage" and "电力故障" in tags:
            boost = 1.15
        if f == "comms_outage" and "通信故障" in tags:
            boost = 1.10
        if f == "supply_shock" and "供应链冲击" in tags:
            boost = 1.20
        if f == "finance_risk" and "金融风险" in tags:
            boost = 1.15
        if f == "panic" and "舆情/恐慌" in tags:
            boost = 0.95  # 舆情通常衰减更快一些
        return max(1.0, half_life * s * boost)

    schedule = {}
    for f in ["supply_shock", "power_outage", "comms_outage", "panic", "finance_risk", "governance_stress"]:
        hl = factor_half_life(f)
        schedule[f] = {t: round(float(0.5 ** (max(0, t - 1) / hl)), 3) for t in range(1, steps + 1)}
    schedule["_meta"] = {
        "duration_hours_est": int(dur_h),
        "step_hours": step_hours,
        "base_half_life_steps": round(float(half_life), 2),
        "factor_half_life_steps": {f: round(float(factor_half_life(f)), 2) for f in ["supply_shock","power_outage","comms_outage","panic","finance_risk","governance_stress"]},
        "tags_seen": tags,
    }
    return schedule

def parse_dt(s: str):
    try:
        s = (s or "").strip()
        if not s:
            return None
        return datetime.strptime(s, "%Y-%m-%d %H:%M")
    except Exception:
        return None

def build_time_axis(steps: int, step_hours: float, hf_params: dict) -> dict:
    """
    为每一步构建时间轴：step -> timestamp/时段标签。
    若提供 start/end，则以 start 为基准；否则仅生成“第N步/第N小时”。
    """
    hp = hf_params or {}
    sh = float(step_hours or 12.0)
    start_dt = parse_dt(hp.get("start_time_str", ""))
    end_dt = parse_dt(hp.get("end_time_str", ""))
    axis = {"start": None, "end": None, "step_hours": sh, "steps": int(steps), "points": {}}

    if start_dt and end_dt and end_dt > start_dt:
        # 若给了起止时间，使用真实跨度来微调 step_hours（保持不删原：只影响时间轴展示与昼夜节律）
        total_h = (end_dt - start_dt).total_seconds() / 3600.0
        sh = max(1.0, total_h / max(1, steps))
        axis["step_hours"] = sh
        axis["start"] = start_dt.strftime("%Y-%m-%d %H:%M")
        axis["end"] = end_dt.strftime("%Y-%m-%d %H:%M")
    elif start_dt:
        axis["start"] = start_dt.strftime("%Y-%m-%d %H:%M")

    for i in range(1, int(steps) + 1):
        if start_dt:
            t = start_dt + pd.Timedelta(hours=sh * (i - 1))
            hour = int(t.hour)
            is_night = bool(hour < 6 or hour >= 20)
            axis["points"][i] = {
                "ts": t.strftime("%Y-%m-%d %H:%M"),
                "hour": hour,
                "is_night": is_night,
            }
        else:
            axis["points"][i] = {"ts": None, "hour": None, "is_night": None}
    return axis

# --- 动态展示名称池：根据事件语义随机选择，体现“像真人分析”的灵活度 ---
RESOURCE_DISPLAY_POOLS = {
    "警力": ["警力", "消防力量", "救援队伍", "安保力量", "武警", "志愿者队伍", "秩序维护力量", "应急响应队", "治安力量", "抢险队伍", "维稳力量", "疏散引导队"],
    "医疗": ["医疗资源", "医护力量", "救护能力", "急救储备", "药品供应", "病床容量", "防疫物资", "医疗机构", "急救车辆", "医护梯队", "医疗保障", "卫生应急"],
    "物资": ["物资储备", "食品供应", "饮用水", "应急物资", "救灾物资", "生活必需品", "粮油储备", "救灾包", "储备粮", "应急装备", "补给物资", "生活保障"],
    "电力": ["电力供应", "能源储备", "供电能力", "燃油储备", "应急电源", "发电能力", "燃料储备", "电力保障", "备用电力", "能源调度", "供电设施", "发电容量"],
    "通信": ["通信保障", "网络覆盖", "信息传播", "广播系统", "联络能力", "通信能力", "互联网接入", "移动网络", "信息基础设施", "通信设施", "联络通道", "信息渠道"],
    "交通": ["交通运力", "物流能力", "运输保障", "道路通行", "配送能力", "交通管制", "疏散通道", "运输网络", "物流通道", "交通保障", "运力储备", "路网状况"],
    "供水": ["供水保障", "自来水供应", "水务调度", "净水能力", "应急供水", "供水管网", "水源保障", "送水能力", "水务应急队", "水厂产能"],
    "燃气": ["燃气供应", "燃气管网", "燃气保障", "燃料供给", "气源调度", "燃气抢修", "供气能力", "燃气储备", "燃气调压站", "燃气应急队"],
    "工程抢修": ["工程抢修力量", "市政抢修队", "线路抢修队", "抢修人员", "应急维修队", "检修力量", "工程保障", "设备检修", "设施抢修", "应急工程队"],
    "财政资金": ["财政资金", "应急资金", "专项拨付", "资金保障", "财政调度", "救助资金", "资金池", "财政支持", "紧急预算", "应急专项资金"],
    "网信/舆情": ["网信力量", "舆情引导", "信息发布能力", "谣言处置力量", "宣传动员", "信息通报", "舆情监测", "应急发布机制", "权威发布", "网络治理力量"],
    "安置保障": ["安置保障", "临时安置点", "住宿保障", "群众安置", "救助保障", "安置资源", "生活安置", "救助站点", "临时收容", "安置服务能力"],
}
INFRA_DISPLAY_POOLS = {
    "电力": ["电力供应网", "供电网络", "电网系统", "能源基础设施", "配电系统", "电力设施", "供电体系", "输电网络", "电网", "供电保障"],
    "通信": ["通信网络", "互联网", "移动网络", "信息基础设施", "通信设施", "网络接入", "通信系统", "信息网络", "通信保障", "网络覆盖"],
    "交通": ["交通路网", "道路系统", "物流通道", "交通基础设施", "运输网络", "疏散通道", "路网体系", "交通设施", "道路通行", "交通保障"],
    "供水": ["供水系统", "自来水网", "饮水保障", "供水基础设施", "水网", "水源供应", "供水设施", "供水管网", "饮水系统", "供水保障"],
}

def _select_display_names(pools: dict, event: str, seed_val: int) -> dict:
    """
    根据事件文本与随机种子，为每个内部键选取展示名称。
    事件关键词匹配的名称权重更高，其余随机；保证可复现（相同 event+seed 得相同结果）。
    """
    event_lower = (event or "").lower()
    rng = random.Random(seed_val)
    result = {}
    for internal_key, names in pools.items():
        scores = []
        for name in names:
            score = 1.0  # 基础随机权重
            # 整词或长词匹配：名称中的 2 字以上片段若出现在事件中则加分
            for i in range(len(name) - 1):
                seg = name[i:i+2]
                if seg in event_lower:
                    score += 1.5
            if len(name) >= 3:
                for i in range(len(name) - 2):
                    seg = name[i:i+3]
                    if seg in event_lower:
                        score += 2.0
            if name in event_lower:
                score += 3.0
            scores.append(max(0.1, score))
        total = sum(scores) + 1e-6
        weights = [s / total for s in scores]
        chosen = rng.choices(names, weights=weights, k=1)[0]
        result[internal_key] = chosen
    return result

def select_resource_display_names(event: str, seed_val: int) -> dict:
    return _select_display_names(RESOURCE_DISPLAY_POOLS, event, seed_val)

def select_infra_display_names(event: str, seed_val: int) -> dict:
    return _select_display_names(INFRA_DISPLAY_POOLS, event, seed_val)

def extract_evidence_factors(facts_text: str) -> dict:
    """
    将检索片段做轻量结构化：仅用于机制模型的参数调制（不是“让 AI 编造”）。
    返回 0..1 的因子：supply_shock, power_outage, comms_outage, panic, finance_risk, governance_stress
    """
    t = (facts_text or "").lower()
    # 关键词触发（可逐步替换成更严谨的抽取器/分类器）
    def has_any(keys):
        return any(k in t for k in keys)

    supply = 0.2 + (0.4 if has_any(["供应链", "运输", "物流", "断供", "shortage", "supply chain"]) else 0)
    power = 0.1 + (0.6 if has_any(["停电", "断电", "电力", "power outage", "blackout"]) else 0)
    comms = 0.1 + (0.5 if has_any(["通信", "断网", "网络中断", "comms", "internet outage"]) else 0)
    panic = 0.1 + (0.6 if has_any(["恐慌", "抢购", "谣言", "panic", "hoard"]) else 0)
    finance = 0.1 + (0.6 if has_any(["挤兑", "bank run", "流动性", "金融危机", "兑付"]) else 0)
    gov = 0.2 + (0.4 if has_any(["应急", "救援", "治理", "维稳", "emergency response"]) else 0)

    return {
        "supply_shock": clamp01(supply),
        "power_outage": clamp01(power),
        "comms_outage": clamp01(comms),
        "panic": clamp01(panic),
        "finance_risk": clamp01(finance),
        "governance_stress": clamp01(gov),
    }

def extract_evidence_factors_from_items(evidence_items: list[dict], fallback_text: str = "") -> tuple[dict, dict]:
    """
    高保真证据映射：
    - 对每条证据（snippet/title）进行关键词匹配，得到各因子 0..1 的贡献分
    - 输出：factors（全局因子 0..1）与 factor_map（每个因子对应的证据贡献列表）
    """
    items = evidence_items or []
    # factor -> keywords (lowercase)
    kw = {
        "supply_shock": ["供应链", "运输", "物流", "断供", "缺货", "shortage", "supply chain", "transport", "logistics"],
        "power_outage": ["停电", "断电", "电力", "blackout", "power outage", "grid failure"],
        "comms_outage": ["通信", "断网", "网络中断", "internet outage", "comms", "telecom", "cellular"],
        "panic": ["恐慌", "抢购", "囤积", "谣言", "panic", "hoard", "rumor"],
        "finance_risk": ["挤兑", "bank run", "流动性", "兑付", "金融危机", "liquidity"],
        "governance_stress": ["应急", "救援", "治理", "维稳", "emergency", "response", "enforcement"],
    }

    factor_map = {k: [] for k in kw.keys()}
    counter_map = {k: [] for k in kw.keys()}
    # 反向（缓解/恢复）关键词：用于冲突检测与净效应计算
    counter_kw = {
        "supply_shock": ["恢复供应", "供给恢复", "补货", "恢复运输", "resumed supply", "restored logistics"],
        "power_outage": ["恢复供电", "电力恢复", "抢修完成", "power restored", "restored power"],
        "comms_outage": ["恢复通信", "网络恢复", "通信恢复", "service restored", "network restored"],
        "panic": ["秩序恢复", "安抚", "澄清", "rumor debunked", "calmed"],
        "finance_risk": ["流动性充足", "挤兑缓解", "稳定金融", "liquidity improved", "stabilized"],
        "governance_stress": ["协调顺畅", "响应有效", "处置到位", "effective response"],
    }
    # 对每条证据打分（0..1）：命中关键词越多越高，并乘以证据强度（strength）
    for it in items:
        text = f"{it.get('title','')} {it.get('snippet','')}".lower()
        eid = it.get("evidence_id", "")
        strength = float(it.get("strength", 0.5))
        for f, keys in kw.items():
            matched = [k for k in keys if k.lower() in text]
            if matched:
                # 轻量权重：命中数量/3，封顶 1
                hit_score = clamp01(len(matched) / 3.0)
                score = clamp01(hit_score * strength)
                factor_map[f].append(
                    {
                        "evidence_id": eid,
                        "score": round(float(score), 2),
                        "matched": matched[:6],
                        "strength": round(float(strength), 2),
                    }
                )
            # 反向证据（缓解/恢复）
            ckeys = counter_kw.get(f, [])
            c_matched = [k for k in ckeys if k.lower() in text]
            if c_matched:
                c_hit = clamp01(len(c_matched) / 2.0)
                c_score = clamp01(c_hit * strength)
                counter_map[f].append(
                    {
                        "evidence_id": eid,
                        "score": round(float(c_score), 2),
                        "matched": c_matched[:6],
                        "strength": round(float(strength), 2),
                    }
                )

    # 聚合全局因子：净支持（support - counter），并给出冲突/置信度
    factors = {}
    confidence = {}
    conflicts = {}
    for f, contribs in factor_map.items():
        top = sorted(contribs, key=lambda x: x["score"], reverse=True)[:3]
        top_c = sorted(counter_map.get(f, []), key=lambda x: x["score"], reverse=True)[:3]
        base = 0.10
        support = sum([float(c["score"]) for c in top])
        counter = sum([float(c["score"]) for c in top_c])
        net = max(0.0, support - 0.8 * counter)
        factors[f] = clamp01(base + net * 0.35)
        # 置信度：支持占比（避免“强反证”时仍给高置信度）
        confidence[f] = round(float(support / (support + counter + 1e-6)), 2)
        conflicts[f] = {
            "support": round(float(support), 2),
            "counter": round(float(counter), 2),
            "has_conflict": bool(counter >= max(0.15, 0.6 * support)),
        }
        factor_map[f] = top
        counter_map[f] = top_c

    # 若没有任何证据条目（或全部未命中），回退到旧的文本抽取逻辑
    if not items or all(len(v) == 0 for v in factor_map.values()):
        factors = extract_evidence_factors(fallback_text or "")
        factor_map = {k: [] for k in factors.keys()}
        counter_map = {k: [] for k in factors.keys()}
        confidence = {k: 0.0 for k in factors.keys()}
        conflicts = {k: {"support": 0.0, "counter": 0.0, "has_conflict": False} for k in factors.keys()}

    factor_map_meta = {
        "top_contributions": factor_map,
        "counter_contributions": counter_map,
        "confidence": confidence,
        "conflicts": conflicts,
    }
    return factors, factor_map_meta

def intervention_profile(name: str) -> dict:
    # 成本（消耗资源）与效果（降低风险/焦虑、提升行政）的机制化参数
    # 资源类型扩展：覆盖水务、燃气、抢修、财政、网信舆情、安置等常见处置资源
    RESOURCE_KEYS = ["警力", "医疗", "物资", "电力", "通信", "交通", "供水", "燃气", "工程抢修", "财政资金", "网信/舆情", "安置保障"]
    profiles = {
        "无": {"cost": {"警力": 1}, "effect": {"calm": 0.00, "order": 0.00, "restore": 0.00}},
        "紧急物资投放": {"cost": {"警力": 2, "物资": 12, "交通": 6, "医疗": 2, "财政资金": 6, "安置保障": 4}, "effect": {"calm": 0.10, "order": 0.03, "restore": 0.04}},
        "媒体安抚": {"cost": {"通信": 4, "警力": 1, "网信/舆情": 6, "财政资金": 2}, "effect": {"calm": 0.14, "order": 0.02, "restore": 0.00}},
        "加强安保": {"cost": {"警力": 10, "交通": 4, "通信": 2}, "effect": {"calm": 0.03, "order": 0.12, "restore": 0.00}},
        "全城宵禁": {"cost": {"警力": 14, "交通": 10, "物资": 3, "通信": 2, "财政资金": 3}, "effect": {"calm": 0.05, "order": 0.18, "restore": 0.00}},
        "电力抢修": {"cost": {"电力": 10, "交通": 4, "物资": 4, "医疗": 1, "工程抢修": 10, "财政资金": 5}, "effect": {"calm": 0.02, "order": 0.04, "restore": 0.16}},
    }
    prof = profiles.get(name, profiles["无"])
    cost = dict((prof.get("cost") or {}))
    for k in RESOURCE_KEYS:
        cost.setdefault(k, 0)
    return {"cost": cost, "effect": prof.get("effect", {"calm": 0.0, "order": 0.0, "restore": 0.0})}

def init_infrastructure(factors: dict) -> dict:
    """
    基础设施子系统（0-100）：电力/通信/交通/供水。
    由证据因子影响初始水平，后续演化会反向影响社会与资源执行效率。
    """
    power = 80 - 45 * factors.get("power_outage", 0.0) + random.uniform(-5, 5)
    comms = 80 - 40 * factors.get("comms_outage", 0.0) + random.uniform(-5, 5)
    transport = 78 - 25 * factors.get("supply_shock", 0.0) + random.uniform(-6, 6)
    water = 82 - 18 * factors.get("supply_shock", 0.0) + random.uniform(-4, 4)
    return {"电力": clamp100(power), "通信": clamp100(comms), "交通": clamp100(transport), "供水": clamp100(water)}

def infra_degrade_and_restore(infra: dict, factors: dict, intervention: str, effect_scale: float) -> dict:
    """
    基础设施演化：冲击会持续侵蚀，特定干预可修复（电力抢修/媒体安抚影响通信等）。
    """
    inf = {k: float(infra.get(k, 80)) for k in ["电力", "通信", "交通", "供水"]}
    # 冲击侵蚀
    inf["电力"] -= 4.0 * factors.get("power_outage", 0.0) + random.uniform(-1.2, 1.2)
    inf["通信"] -= 3.5 * factors.get("comms_outage", 0.0) + random.uniform(-1.0, 1.0)
    inf["交通"] -= 2.5 * factors.get("supply_shock", 0.0) + random.uniform(-1.5, 1.5)
    inf["供水"] -= 1.8 * factors.get("supply_shock", 0.0) + random.uniform(-1.0, 1.0)

    # 干预修复
    if intervention == "电力抢修":
        inf["电力"] += 10.0 * effect_scale + random.uniform(-1.0, 1.0)
        inf["通信"] += 2.0 * effect_scale
    if intervention == "媒体安抚":
        inf["通信"] += 5.0 * effect_scale + random.uniform(-0.8, 0.8)
    if intervention == "全城宵禁":
        inf["交通"] -= 6.0  # 宵禁直接抑制交通可用性

    return {k: clamp100(v) for k, v in inf.items()}

def mechanistic_step(state: dict, resources: dict, factors: dict, intervention: str, step: int, infra: dict, params: dict, evidence_items: list[dict], factor_map: dict, decay_schedule: dict, time_axis: dict) -> tuple[dict, dict, dict, dict]:
    """
    机制化演化：状态由可审计规则决定（资源约束、冲击因子、干预成本/效果、阈值触发）
    返回 (new_state, new_resources, new_infra, audit)
    """
    s = {k: float(state.get(k, 50)) for k in ["行政效能", "焦虑指数", "资源缺口", "动荡风险"]}
    r = {k: float(resources.get(k, 50)) for k in ["警力", "医疗", "物资", "电力", "通信", "交通", "供水", "燃气", "工程抢修", "财政资金", "网信/舆情", "安置保障"]}
    inf = {k: float(infra.get(k, 80)) for k in ["电力", "通信", "交通", "供水"]}

    prof = intervention_profile(intervention)
    cost = prof["cost"]
    eff = prof["effect"]

    # 证据引用：使用 factor_map 的 top 贡献证据作为可追溯引用
    cited = []
    top_map = (factor_map or {}).get("top_contributions", {}) if isinstance(factor_map, dict) else {}
    for f, top in (top_map or {}).items():
        for c in (top or []):
            if isinstance(c, dict) and c.get("evidence_id"):
                cited.append(c["evidence_id"])
    cited = sorted(list({c for c in cited if c}))

    # 资源可用性（不足时效果打折 + 产生“执行失败”审计项）
    shortage_ratio = 0.0
    for k, c in cost.items():
        avail = r.get(k, 0.0)
        if c > 0:
            shortage_ratio = max(shortage_ratio, max(0.0, (c - avail) / c))

    execution_penalty = clamp01(shortage_ratio)
    # 可行性硬约束：基础设施与关键资源不足时，判定“不可执行”，不扣资源、不产生效果
    min_infra = float(params.get("min_infra_to_execute", 25))
    infra_ok = min(inf["电力"], inf["通信"], inf["交通"]) >= min_infra
    resources_ok = execution_penalty <= 0.15
    executable = bool(infra_ok and resources_ok)
    effect_scale = (1.0 - 0.7 * execution_penalty) if executable else 0.0

    # 昼夜节律：夜间更易焦虑/动荡，且执行折扣（不改变原结构，只对高保真新增）
    is_night = False
    if isinstance(time_axis, dict):
        p = (time_axis.get("points") or {}).get(step) or {}
        is_night = bool(p.get("is_night")) if p.get("is_night") is not None else False
    if bool(params.get("circadian", True)) and is_night:
        effect_scale = max(0.0, effect_scale * (1.0 - float(params.get("night_exec_penalty", 0.15))))

    # 先扣资源（不低于 0）
    if executable:
        for k, c in cost.items():
            r[k] = max(0.0, r.get(k, 0.0) - float(c))

    # 冲击：由证据因子调制（越像真实案例/危机类型，对应维度冲击更强）
    ds = decay_schedule or {}
    d = {
        "panic": float(((ds.get("panic") or {}).get(step, 1.0))),
        "comms_outage": float(((ds.get("comms_outage") or {}).get(step, 1.0))),
        "finance_risk": float(((ds.get("finance_risk") or {}).get(step, 1.0))),
        "supply_shock": float(((ds.get("supply_shock") or {}).get(step, 1.0))),
        "power_outage": float(((ds.get("power_outage") or {}).get(step, 1.0))),
        "governance_stress": float(((ds.get("governance_stress") or {}).get(step, 1.0))),
    }
    shock_anxiety = params.get("shock_anxiety_k", 1.0) * (
        8.0 * factors["panic"] * d["panic"]
        + 6.0 * factors["comms_outage"] * d["comms_outage"]
        + 5.0 * factors["finance_risk"] * d["finance_risk"]
    )
    shock_gap = params.get("shock_gap_k", 1.0) * (
        9.0 * factors["supply_shock"] * d["supply_shock"]
        + 7.0 * factors["power_outage"] * d["power_outage"]
    )
    shock_risk = params.get("shock_risk_k", 1.0) * (
        7.0 * factors["panic"] * d["panic"]
        + 8.0 * factors["supply_shock"] * d["supply_shock"]
        + 6.0 * factors["finance_risk"] * d["finance_risk"]
    )
    shock_admin = -params.get("shock_admin_k", 1.0) * (
        6.0 * factors["power_outage"] * d["power_outage"]
        + 4.0 * factors["comms_outage"] * d["comms_outage"]
        + 5.0 * factors["supply_shock"] * d["supply_shock"]
    )

    # 干预效果：降低焦虑/动荡，提升行政，缓解缺口（restore 主要影响电力/通信带来的行政恢复）
    calm = 12.0 * eff["calm"] * effect_scale
    order = 12.0 * eff["order"] * effect_scale
    restore = 12.0 * eff["restore"] * effect_scale

    # 动力学耦合：缺口↑ 会推高焦虑与风险；行政↓ 会推高风险
    coupling_anxiety = params.get("coupling_anxiety_k", 1.0) * (0.10 * (s["资源缺口"] - 50.0) + 0.08 * (50.0 - s["行政效能"]))
    coupling_risk = params.get("coupling_risk_k", 1.0) * (0.12 * (s["焦虑指数"] - 50.0) + 0.10 * (s["资源缺口"] - 50.0) + 0.10 * (50.0 - s["行政效能"]))

    # 基础设施对行政与缺口的反馈：电力/交通越差，行政越难、缺口越大；通信差会推高焦虑
    infra_admin_pen = (80.0 - inf["电力"]) * 0.06 + (80.0 - inf["交通"]) * 0.04
    infra_gap_pen = (80.0 - inf["交通"]) * 0.06 + (80.0 - inf["电力"]) * 0.05
    infra_anx_pen = (80.0 - inf["通信"]) * 0.05

    # 宵禁副作用：焦虑上升、交通受限 -> 缺口可能加剧
    curfew_penalty = 0.0
    if intervention == "全城宵禁":
        curfew_penalty = 6.0
        s["资源缺口"] += 4.0

    # 计算新状态（加入少量噪声但可复现：由外部 seed 控制 random）
    s2 = {}
    s2["行政效能"] = s["行政效能"] + shock_admin + restore - infra_admin_pen + random.uniform(-1.5, 1.5)
    s2["资源缺口"] = s["资源缺口"] + shock_gap - 6.0 * eff["restore"] * effect_scale + infra_gap_pen + random.uniform(-1.5, 1.5)
    anxiety_boost = float(params.get("night_anxiety_boost", 0.25)) if (bool(params.get("circadian", True)) and is_night) else 0.0
    risk_boost = float(params.get("night_risk_boost", 0.20)) if (bool(params.get("circadian", True)) and is_night) else 0.0
    s2["焦虑指数"] = s["焦虑指数"] + shock_anxiety * (1.0 + anxiety_boost) - calm + coupling_anxiety + infra_anx_pen + curfew_penalty + random.uniform(-2.0, 2.0)
    s2["动荡风险"] = s["动荡风险"] + shock_risk * (1.0 + risk_boost) - order + coupling_risk + random.uniform(-2.0, 2.0)

    # 更新基础设施（冲击持续 + 干预修复）
    new_infra = infra_degrade_and_restore(infra, factors, intervention, effect_scale)

    # 阈值触发：风险过高触发“骚乱/踩踏”事件，消耗警力医疗并进一步推高焦虑
    triggered = []
    if s2["动荡风险"] >= float(params.get("trigger_risk", 80)):
        triggered.append("高动荡阈值触发：局部骚乱/聚集事件")
        r["警力"] = max(0.0, r["警力"] - 8.0)
        r["医疗"] = max(0.0, r["医疗"] - 4.0)
        s2["焦虑指数"] += 4.0
    if s2["资源缺口"] >= float(params.get("trigger_gap", 75)):
        triggered.append("高缺口阈值触发：排队与争抢加剧")
        s2["焦虑指数"] += 3.0
        s2["动荡风险"] += 2.0
    if s2["行政效能"] <= float(params.get("trigger_admin", 25)):
        triggered.append("低行政阈值触发：协调失灵与执行落差")
        s2["动荡风险"] += 3.0

    new_state = {k: clamp100(v) for k, v in s2.items()}
    new_resources = {k: clamp100(v) for k, v in r.items()}

    audit = {
        "step": step,
        "intervention": intervention,
        "executable": executable,
        "resource_shortage_ratio": round(execution_penalty, 2),
        "effect_scale": round(effect_scale, 2),
        "triggered": triggered,
        "resource_cost": cost,
        "factors": {k: round(float(v), 2) for k, v in factors.items()},
        "infra": {k: clamp100(v) for k, v in inf.items()},
        "infra_after": new_infra,
        "evidence_citations": cited,
        "factor_contributions": (factor_map or {}).get("top_contributions", factor_map) if isinstance(factor_map, dict) else (factor_map or {}),
        "factor_confidence": (factor_map or {}).get("confidence", {}) if isinstance(factor_map, dict) else {},
        "factor_conflicts": (factor_map or {}).get("conflicts", {}) if isinstance(factor_map, dict) else {},
        "decay_meta": (decay_schedule or {}).get("_meta", {}),
        "time_point": ((time_axis or {}).get("points") or {}).get(step, {}),
    }
    return new_state, new_resources, new_infra, audit

def narrate_from_mechanism(event: str, facts: str, step: int, intervention: str, state: dict, prev_state: dict, resources: dict, audit: dict) -> dict:
    """
    将“机制计算结果”转成四方叙事。AI 只能基于输入的状态/审计，不允许凭空造数据。
    """
    delta = {k: int(state[k]) - int(prev_state.get(k, 0)) for k in ["行政效能", "焦虑指数", "资源缺口", "动荡风险"]}
    base_context = f"""事件：{event}
阶段：第{step}步
干预：{intervention}
状态（0-100）：{state}
变化量：{delta}
资源（0-100）：{resources}
审计要点：可执行={audit.get('executable')}；资源不足比例={audit.get('resource_shortage_ratio')}；触发事件={audit.get('triggered')}；证据引用={audit.get('evidence_citations')}
实证摘要（仅供背景，不要求逐字引用）：{(facts or '')[:600]}
"""
    if client is None:
        # 无模型时使用规则化模板（仍然基于机制结果）
        return {
            "official": f"第{step}阶段，官方尝试执行「{intervention}」（可执行={audit.get('executable')}）。行政效能{delta['行政效能']:+d}，资源缺口{delta['资源缺口']:+d}。重点风险为动荡风险{state['动荡风险']}%，需围绕资源与秩序联动处置。",
            "citizen": f"民众焦虑指数{state['焦虑指数']}%（变化{delta['焦虑指数']:+d}），对物资与信息的可获得性更敏感，出现排队/囤积倾向的概率上升。",
            "media": f"媒体报道聚焦指标变化与触发事件：{'; '.join(audit.get('triggered') or ['无'])}。舆情可能放大焦虑并影响社会预期。",
            "audit": f"审计结论：可执行={audit.get('executable')}；资源不足比例={audit.get('resource_shortage_ratio')}；效果缩放={audit.get('effect_scale')}；证据引用={audit.get('evidence_citations')}。建议优先保障关键资源并避免指标与叙事不一致。"
        }

    prompt = f"""你是推演系统的“叙事生成器”，必须严格遵守审计约束：
1) 不得编造任何未在输入中出现的数字、事件、资源水平；
2) 所有结论必须可追溯到状态/变化量/审计要点；
3) 输出必须是 JSON，键为 official/citizen/media/audit，值为中文文本，每段 120-220 字。

输入如下：
{base_context}
"""
    try:
        res = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            timeout=30
        )
        j = safe_json(res.choices[0].message.content)
        return {
            "official": str(j.get("official", ""))[:1200],
            "citizen": str(j.get("citizen", ""))[:1200],
            "media": str(j.get("media", ""))[:1200],
            "audit": str(j.get("audit", ""))[:1200],
        }
    except Exception as e:
        return {
            "official": f"[叙事生成异常] {str(e)[:80]}",
            "citizen": "叙事生成失败（请检查 API）。",
            "media": "叙事生成失败（请检查 API）。",
            "audit": "叙事生成失败（请检查 API）。"
        }

def get_evidence(q):
    if not serper_key:
        return "离线仿真模式"
    try:
        r = requests.post("https://google.serper.dev/search",
            headers={"X-API-KEY": serper_key, "Content-Type": "application/json"},
            json={"q": f"{q} 社会演化 历史案例 危机应对 PESTEL", "num": 8}, timeout=12)
        return "\n".join([x.get("snippet","") for x in r.json().get("organic", [])])
    except:
        return "实证检索失败，使用内部模型"

def get_evidence_items(q: str) -> tuple[str, list[dict]]:
    """
    证据条目化：为高保真模式提供可引用 evidence_id。
    返回 (facts_text, items)。facts_text 仍然兼容原有流程。
    """
    if not serper_key:
        return "离线仿真模式", []
    try:
        r = requests.post(
            "https://google.serper.dev/search",
            headers={"X-API-KEY": serper_key, "Content-Type": "application/json"},
            json={"q": f"{q} 社会演化 历史案例 危机应对 PESTEL", "num": 8},
            timeout=12,
        )
        organic = r.json().get("organic", []) or []
        items = []
        for idx, x in enumerate(organic, start=1):
            base = {
                "evidence_id": f"E{idx:02d}",
                "title": x.get("title", ""),
                "link": x.get("link", ""),
                "snippet": x.get("snippet", ""),
            }
            base.update(score_evidence_strength(base))
            base.update(extract_evidence_struct(base))
            items.append(base)
        facts_text = "\n".join([it.get("snippet", "") for it in items if it.get("snippet")])
        return facts_text, items
    except Exception:
        return "实证检索失败，使用内部模型", []

def _default_matrix():
    return {"行政效能": 50, "焦虑指数": 50, "资源缺口": 50, "动荡风险": 50}

def init_state(event, facts):
    ctx = ""
    prompt = f"{ctx}事件：{event}\n事实：{facts}\n返回JSON（0-100）：行政效能、焦虑指数、资源缺口、动荡风险"
    try:
        res = client.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":prompt}], response_format={"type":"json_object"})
        j = safe_json(res.choices[0].message.content)
        return {k: max(0, min(100, int(j.get(k, 50)))) for k in ["行政效能", "焦虑指数", "资源缺口", "动荡风险"]}
    except Exception:
        return _default_matrix()

def init_resources():
    return {
        "警力": random.randint(60,95), "医疗": random.randint(50,90),
        "物资": random.randint(30,85), "电力": random.randint(20,70),
        "通信": random.randint(40,90), "交通": random.randint(30,80),
        "供水": random.randint(40,90), "燃气": random.randint(40,90),
        "工程抢修": random.randint(30,85), "财政资金": random.randint(40,95),
        "网信/舆情": random.randint(35,90), "安置保障": random.randint(30,85),
    }

def evolve_step(event, facts, matrix, step, intervention, resources):
    prompt = f"""
事件：{event}
实证：{facts}
社会状态：{matrix}
资源状态：{resources}
干预措施：{intervention}
第{step}阶段演化。
严格返回JSON：official、citizen、media、audit（详细专业）
"""
    try:
        res = client.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":prompt}], response_format={"type":"json_object"}, timeout=30)
        return safe_json(res.choices[0].message.content)
    except Exception as e:
        return {
            "official": f"[API 异常] 官方应对：{str(e)[:80]}",
            "citizen": "[API 异常] 民众反应数据暂不可用",
            "media": "[API 异常] 媒体视角数据暂不可用",
            "audit": "[API 异常] 审计结论暂不可用"
        }

def update_state(old, evo, intervention):
    adj = {
        "无": (-3,5,-2,6),
        "紧急物资投放": (-2,-5,10,-4),
        "媒体安抚": (-1,-10,2,-3),
        "加强安保": (4,2,-5,-8),
        "全城宵禁": (6,-8,-3,-10),
        "电力抢修": (8,0,-5,2)
    }
    a,b,c,d = adj.get(intervention, adj["无"])
    base = _default_matrix()
    for k in base:
        base[k] = old.get(k, 50)
    return {
        "行政效能": max(0, min(100, base["行政效能"] + a + random.randint(-2,3))),
        "焦虑指数": max(0, min(100, base["焦虑指数"] + b + random.randint(-4,6))),
        "资源缺口": max(0, min(100, base["资源缺口"] + c + random.randint(-3,5))),
        "动荡风险": max(0, min(100, base["动荡风险"] + d + random.randint(-3,7)))
    }

def update_resources(res):
    return {
        k: max(0, min(100, v + random.randint(-8,6)))
        for k,v in res.items()
    }

def gen_black_swan(event):
    if random.random() < 0.3:
        return "无黑天鹅事件"
    try:
        prompt = f"{event} 背景下，生成低概率高冲击突发事件"
        return client.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":prompt}]).choices[0].message.content
    except Exception as e:
        return f"[生成异常] 黑天鹅模块暂不可用：{str(e)[:60]}"

def gen_causal(event):
    try:
        prompt = f"为 {event} 生成PESTEL全维度因果链，结构化输出"
        return client.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":prompt}]).choices[0].message.content
    except Exception as e:
        return f"[因果链生成异常] {str(e)[:100]}\n请检查 API 配置后重新运行。"

def _mp_call_json(role_name: str, prompt: str, timeout: int = 45) -> dict:
    """多视角研判调用：强制 JSON 输出并容错解析。"""
    if not client:
        return {"error": "no_client", "role": role_name}
    try:
        res = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            timeout=timeout,
        )
        j = safe_json(res.choices[0].message.content)
        if not isinstance(j, dict):
            j = {"raw": str(j)}
        j.setdefault("role", role_name)
        return j
    except Exception as e:
        return {"role": role_name, "error": str(e)[:160]}

def multi_perspective_assess(event: str, empirical: str, timeline, swan: str, matrix: dict, resources: dict) -> dict:
    """
    4 视角 + 1 裁决的综合研判：每个视角输出 judgement + forecast + confidence + assumptions + triggers + recommendations。
    最后裁决视角输出综合结论与情景区间。
    """
    ctx = f"""事件：{event}
实证材料：{(empirical or '')[:12000]}
演化时序：{str(timeline)[:6000]}
黑天鹅：{swan}
社会状态（0-100）：{matrix}
资源状态（0-100）：{resources}
"""
    common_rules = """你必须遵守：
1) 不得编造任何未在输入中出现的事实或数字；
2) 若信息不足必须明确写“不确定/需补充”；
3) 必须输出 JSON，包含字段：judgement（判断）、forecast（推演预测）、confidence（0-1）、assumptions（要点列表）、triggers（触发条件列表）、recommendations（建议列表）。
4) judgement 与 forecast 必须分开写清楚，且每段 120-220 字。
"""

    prompts = {
        "保守派风险官": f"""你是“保守派风险官”，目标是避免过度反应，强调证据与稳健。
{common_rules}
输入如下：
{ctx}
""",
        "危机派应急官": f"""你是“危机派应急官”，目标是识别最坏情境与脆弱环节，提前布防，但仍不得编造。
{common_rules}
输入如下：
{ctx}
""",
        "机制/数据审计官": f"""你是“机制/数据审计官”，只做一致性审计与逻辑校验：哪些推断站得住，哪些是过推；指出矛盾与缺口。
{common_rules}
输入如下：
{ctx}
""",
        "民生/社会心理官": f"""你是“民生/社会心理官”，评估公众行为、舆情传播、群体事件的现实边界与反馈回路，给出更贴近实际的预判。
{common_rules}
输入如下：
{ctx}
""",
    }

    views = {}
    for role_name, p in prompts.items():
        views[role_name] = _mp_call_json(role_name, p, timeout=45)

    # 裁决：汇总分歧、给出情景区间与最终综合建议
    arb_rules = """你是“裁决/汇总官”。你必须：
1) 逐条列出四个视角的主要分歧点（至少3条）；
2) 给出综合 judgement（判断）与综合 forecast（推演预测），并输出三情景：保守/基准/悲观（每个 80-140 字）；
3) 给出综合 confidence（0-1），以及最关键的 5 条 triggers（触发条件）与 6 条 recommendations（建议）。
4) 严禁编造输入中没有的事实或数字。
严格返回 JSON，字段：disagreements（列表）、judgement、forecast、scenarios（对象含 conservative/base/bear）、confidence、triggers（列表）、recommendations（列表）。
"""
    arb_prompt = f"""{arb_rules}
输入上下文：
{ctx}
四视角输出（JSON）：
{json.dumps(views, ensure_ascii=False)[:12000]}
"""
    arbitration = _mp_call_json("裁决/汇总官", arb_prompt, timeout=60)
    return {"views": views, "arbitration": arbitration}

def _format_mp_block(mp: dict) -> str:
    """将多视角研判结果格式化为可读 Markdown。"""
    if not mp or not isinstance(mp, dict):
        return ""
    views = mp.get("views") or {}
    arb = mp.get("arbitration") or {}

    def _fmt_list(xs):
        if not xs:
            return "—"
        if isinstance(xs, str):
            return xs
        if isinstance(xs, list):
            return "\n".join([f"- {str(x)[:220]}" for x in xs[:12]]) or "—"
        return str(xs)[:600]

    out = "## 九、多视角研判（4 视角 + 1 裁决）\n\n"
    # 4 views
    for role_name in ["保守派风险官", "危机派应急官", "机制/数据审计官", "民生/社会心理官"]:
        v = views.get(role_name) or {}
        out += f"### 视角：{role_name}\n\n"
        if v.get("error"):
            out += f"- 调用异常：{v.get('error')}\n\n"
            continue
        out += f"- judgement：{str(v.get('judgement','')).strip()}\n"
        out += f"- forecast：{str(v.get('forecast','')).strip()}\n"
        out += f"- confidence：{v.get('confidence','')}\n"
        out += f"- assumptions：\n{_fmt_list(v.get('assumptions'))}\n"
        out += f"- triggers：\n{_fmt_list(v.get('triggers'))}\n"
        out += f"- recommendations：\n{_fmt_list(v.get('recommendations'))}\n\n"

    # arbitration
    out += "### 裁决：综合结论（汇总官）\n\n"
    if arb.get("error"):
        out += f"- 调用异常：{arb.get('error')}\n"
        return out
    out += f"- disagreements：\n{_fmt_list(arb.get('disagreements'))}\n"
    out += f"- judgement：{str(arb.get('judgement','')).strip()}\n"
    out += f"- forecast：{str(arb.get('forecast','')).strip()}\n"
    out += f"- scenarios：\n"
    sc = arb.get("scenarios") or {}
    if isinstance(sc, dict):
        out += f"  - conservative：{str(sc.get('conservative','')).strip()}\n"
        out += f"  - base：{str(sc.get('base','')).strip()}\n"
        out += f"  - bear：{str(sc.get('bear','')).strip()}\n"
    else:
        out += f"  - {str(sc)[:600]}\n"
    out += f"- confidence：{arb.get('confidence','')}\n"
    out += f"- triggers：\n{_fmt_list(arb.get('triggers'))}\n"
    out += f"- recommendations：\n{_fmt_list(arb.get('recommendations'))}\n\n"
    return out

def gen_report(event, facts, timeline, swan, matrix, resources):
    empirical = facts
    prompt = f"""
你是国家级战略安全顾问，生成正式研判报告。

事件：{event}
实证材料：{empirical}
演化时序：{timeline}
黑天鹅：{swan}
社会状态：{matrix}
资源状态：{resources}

若实证材料中包含【用户上传的种子材料】，请直接从该材料全文理解分析对象（企业名、年度、文档类型等），并据此撰写报告，不得臆造。若无上传材料则基于事件与演化结果研判。
报告必须包含：1. 事件定级 2. 物理瘫痪分析 3. 社会临界点 4. 演化复盘 5. 风险预警 6. 干预建议 7. 结论
"""
    try:
        base_report = client.chat.completions.create(model="deepseek-chat", messages=[{"role":"user","content":prompt}], timeout=60).choices[0].message.content
        # 直接启用多视角研判，并合并进最终综合研判报告（不加开关）
        mp = multi_perspective_assess(event, empirical, timeline, swan, matrix, resources)
        mp_block = _format_mp_block(mp)
        if mp_block:
            return f"{base_report}\n\n---\n\n{mp_block}"
        return base_report
    except Exception as e:
        return f"【报告生成异常】API 调用失败：{str(e)[:120]}\n请检查网络与 API 配置后重新运行仿真并导出报告。"

def classify_event_relevance(event: str) -> dict:
    """
    智能化判断用户输入是否与政府政策、社会资源、公共危机相关。
    若无关（如续写小说、一般创作），则不应启动危机推演模块。
    返回 {"is_crisis_simulation": bool, "reason": str}
    """
    if not client:
        return {"is_crisis_simulation": True, "reason": "无API时默认启用危机推演"}
    prompt = """你是一个意图分类器。判断用户的输入是否适合进行「政府/社会危机推演仿真」。

以下情况应返回 is_crisis_simulation=true：
- 涉及政府应对、公共政策、社会资源调度
- 涉及基础设施故障（停电/断水/通信/交通中断）
- 涉及金融挤兑、供应链中断、民生保障
- 涉及治安、秩序、救援、应急响应
- 涉及民众恐慌、抢购、动荡等社会演化

以下情况应返回 is_crisis_simulation=false：
- 纯文学创作（续写小说、写诗、写故事）
- 预测虚构作品的情节走向
- 与政府/社会资源无关的创作、问答

用户输入：「{event}」

严格返回JSON，仅包含两个键：is_crisis_simulation（布尔）、reason（一句话说明原因，中文）""".format(event=(event or "")[:500])
    try:
        res = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            timeout=15,
        )
        j = safe_json(res.choices[0].message.content)
        is_crisis = bool(j.get("is_crisis_simulation", True))
        reason = str(j.get("reason", ""))[:200] or ("适合危机推演" if is_crisis else "与政府社会资源无关，采用创意模式")
        return {"is_crisis_simulation": is_crisis, "reason": reason}
    except Exception as e:
        return {"is_crisis_simulation": True, "reason": f"分类异常，默认启用危机推演：{str(e)[:60]}"}

def run_branch_from_step(fork_step: int, new_intervention: str, event: str, facts: str) -> dict:
    """从第 fork_step 步之后用新干预策略继续推演，返回 branch 数据。"""
    timeline = st.session_state.get("timeline") or []
    mh = st.session_state.get("matrix_history") or []
    rh = st.session_state.get("resources") or []
    ih = st.session_state.get("infra_history") or []
    idx = min(fork_step + 1, len(mh) - 1) if mh else 0
    matrix = dict(mh[idx]) if idx < len(mh) and mh else {"行政效能":50,"焦虑指数":50,"资源缺口":50,"动荡风险":50}
    resources = dict(rh[idx]) if idx < len(rh) and rh else {}
    infra = dict(ih[idx]) if idx < len(ih) and ih else {"电力":80,"通信":80,"交通":80,"供水":80}
    factors = st.session_state.get("init_factors") or st.session_state.get("stepping_factors")
    if not factors or (isinstance(factors, dict) and "top_contributions" in factors):
        factors = extract_evidence_factors(facts[:2000])
    eff = int(st.session_state.get("stepping_effective_steps") or len(timeline))
    branch_timeline = []
    branch_matrix_hist = list(st.session_state.matrix_history[:fork_step+2])  # init + steps 1..fork_step+1
    branch_resources = list(st.session_state.resources[:fork_step+2])
    branch_infra = list((st.session_state.get("infra_history") or [])[:fork_step+2])
    for i in range(fork_step + 2, eff + 1):  # 从 step fork_step+2 开始（即分支后的第一步）
        if sim_mode.startswith("高保真"):
            prev = matrix
            matrix, resources, infra, audit = mechanistic_step(
                matrix, resources, factors, new_intervention, i, infra,
                st.session_state.hf_params, st.session_state.evidence_items,
                st.session_state.evidence_factor_map, st.session_state.decay_schedule,
                st.session_state.time_axis,
            )
            evo = narrate_from_mechanism(event, facts, i, new_intervention, matrix, prev, resources, audit)
            branch_timeline.append({"step": i, "data": evo, "audit": audit, "state": matrix.copy(), "resources": resources.copy(), "infra": infra.copy()})
        else:
            evo = evolve_step(event, facts, matrix, i, new_intervention, resources)
            branch_timeline.append({"step": i, "data": evo})
            matrix = update_state(matrix, evo, new_intervention)
            resources = update_resources(resources)
        branch_matrix_hist.append(matrix.copy())
        branch_resources.append(resources.copy())
        if sim_mode.startswith("高保真"):
            branch_infra.append(infra.copy())
    full_tl = list(timeline[:fork_step+1]) + branch_timeline
    report = gen_report(event, facts, full_tl, st.session_state.get("swan",""), matrix, resources)
    return {"fork_step": fork_step, "intervention": new_intervention, "timeline": branch_timeline, "matrix_history": branch_matrix_hist, "resources": branch_resources, "infra_history": branch_infra, "report": report, "full_timeline": full_tl}

def extract_entity_graph(event: str, facts: str) -> dict:
    """从事件与实证中抽取实体与关系，构建简单知识图。"""
    if not client:
        return {"entities": [], "relations": []}
    text = f"事件：{event}\n\n实证/背景：{(facts or '')[:3000]}"
    prompt = f"""从以下文本中抽取关键实体（人物、组织、地点、资源、事件类型等）及其关系。

文本：
{text}

严格返回JSON，格式：
{{"entities": [{{"name": "实体名", "type": "人物|组织|地点|资源|事件"}}], "relations": [{{"source": "实体1", "target": "实体2", "relation": "关系描述"}}]}}

每类实体最多5个，关系最多8条。只输出JSON，无其他内容。"""
    try:
        res = client.chat.completions.create(model="deepseek-chat", messages=[{"role": "user", "content": prompt}], response_format={"type": "json_object"}, timeout=15)
        j = safe_json(res.choices[0].message.content)
        return {"entities": j.get("entities", [])[:15], "relations": j.get("relations", [])[:12]}
    except Exception:
        return {"entities": [], "relations": []}

HISTORY_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "shence_history.json")

def _load_history_from_file() -> list:
    try:
        if os.path.exists(HISTORY_FILE):
            with open(HISTORY_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return []

def _save_to_history(scenario_name: str, event: str, report: str, matrix_history: list):
    try:
        rec = {"name": scenario_name or "未命名", "event": (event or "")[:200], "report_excerpt": (report or "")[:500], "final_state": matrix_history[-1] if matrix_history else {}, "ts": datetime.now().isoformat()}
        hist = _load_history_from_file()
        hist.append(rec)
        with open(HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(hist[-50:], f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def chat_with_perspective(perspective: str, perspective_narrative: str, user_question: str, context: str) -> str:
    """与某一视角（官方/民众/媒体/审计）对话，基于其叙事与推演上下文回答。"""
    if not client:
        return "[需要配置 API] 请配置 API Key 后使用与视角对话功能。"
    role_map = {"官方": "官方决策者", "民众": "普通民众代表", "媒体": "媒体评论员", "审计": "独立审计专家"}
    role = role_map.get(perspective, perspective)
    prompt = f"""你在扮演推演系统中的「{role}」视角。以下是该视角在推演中的叙事与立场摘要：

{perspective_narrative[:800]}

推演背景摘要：
{context[:600]}

用户提问：{user_question}

请以该视角的身份和立场，用第一人称或该角色口吻简洁回答（150字以内）。不要跳出角色，不要以"作为AI"等表述开头。"""
    try:
        return client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            timeout=20,
        ).choices[0].message.content
    except Exception as e:
        return f"[对话异常] {str(e)[:80]}"

def gen_creative_response(event: str) -> str:
    """
    当用户输入与政府/社会危机无关时，直接按用户意图完成任务（如续写、预测走向等）。
    """
    if not client:
        return f"[需要配置 API] 您的请求：{event}\n\n系统判断该请求与政府社会危机推演无关，应直接完成您的创作或预测需求。请配置 API 后重试。"
    prompt = f"""用户的请求与政府政策、社会资源、公共危机无关，请不要进行危机推演分析。

用户请求：{event}

请直接完成用户的请求。例如：
- 若用户要求续写某作品：请按原著风格续写
- 若用户要求预测故事走向：请给出合理的剧情预测
- 若用户是其他创作类请求：请按用户意图完成

输出应直接是用户所需的内容，无需额外说明或框架。"""
    try:
        return client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            timeout=90,
        ).choices[0].message.content
    except Exception as e:
        return f"[生成异常] {str(e)[:120]}\n请检查 API 配置后重试。"

def export_pdf(report: str, doc_appendix: str = ""):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    try:
        pdf.multi_cell(0, 10, report)
    except Exception:
        pdf.multi_cell(0, 10, report.encode("latin-1", errors="replace").decode("latin-1"))
    if doc_appendix and not doc_appendix.startswith("["):
        pdf.add_page()
        pdf.set_font("Arial", size=11, style="B")
        pdf.multi_cell(0, 8, "Appendix: Seed Document Key Points / 种子材料要点")
        pdf.set_font("Arial", size=10)
        try:
            pdf.multi_cell(0, 6, doc_appendix[:4000])
        except Exception:
            pdf.multi_cell(0, 6, doc_appendix[:4000].encode("latin-1", errors="replace").decode("latin-1"))
    return bytes(pdf.output(dest='S'))

# ---------------------- 主界面 ----------------------
_dbg("14-进入主界面")
st.title("事件推演仿真系统")

# 文本种子输入（替代上传，直接粘贴全文用于研判）
st.subheader("📝 文本种子输入（用于综合分析研判）")
seed_text = st.text_area(
    "在此粘贴需要研判的完整文本内容（可来自 PDF、Word、报告等）",
    value=st.session_state.get("uploaded_doc_text", ""),
    height=220,
    placeholder="将需要分析研判的全文粘贴到这里，例如年报、专项报告、政策文件、调查材料等。\n系统会直接基于这些文字做内容分析与最终综合研判。",
)
if seed_text:
    st.session_state.uploaded_doc_text = seed_text
    # 基于全文内容进行一次固定规则的分析与研判
    analysis = analyze_uploaded_text(seed_text)
    st.session_state.uploaded_doc_analysis = analysis
    st.session_state.uploaded_doc_report = generate_uploaded_doc_report(analysis, "用户输入文本", "txt")

_dbg("15-跳过历史引用功能")

# 事件输入（自主输入，无需模板）
st.subheader("📡 初始扰动事件")
event = st.text_area(
    "请自主描述您要推演的事件场景",
    value=st.session_state.get("event", ""),
    height=120,
    placeholder="例如：一线城市核心区域停电停水 48 小时，通信不稳，民众恐慌抢购；或：粮食供应链中断 7 天，米面油库存下降，物价上涨……\n\n也可输入与政府社会无关的请求（如：请续写红楼梦、预测某故事走向），系统将智能识别并直接完成，不启动危机推演模块。"
)

# 控制面板
with st.expander("⚙️ 仿真控制面板", expanded=True):
    c1,c2,c3,c4 = st.columns(4)
    with c1:
        steps = st.slider("演化步数", 1, 8, 4)
    with c2:
        st.caption("推演方式：输入事件后自动全流程演化")
    with c3:
        intervention = st.selectbox("🎯 官方干预策略", [
            "无","紧急物资投放","媒体安抚","加强安保","全城宵禁","电力抢修"
        ], help="无/物资/媒体/安保/宵禁/电力，不同策略会改变社会矩阵演化系数")
    with c4:
        enable_swan = st.checkbox("启用黑天鹅", value=True)
    with st.expander("📖 干预策略说明", expanded=False):
        st.markdown("""
- **无**：无专项干预，按自然演化
- **紧急物资投放**：缓解资源缺口与焦虑，可能略增行政压力
- **媒体安抚**：显著降焦虑，轻微影响行政与资源
- **加强安保**：提升行政效能、压制动荡，略增资源消耗
- **全城宵禁**：强控动荡与行政，但会推高焦虑
- **电力抢修**：优先恢复行政与秩序，缓解资源与动荡
        """)
        # 干预策略成本与效果可视化
        st.markdown("**📊 干预策略成本与效果对比**")
        all_interventions = ["无","紧急物资投放","媒体安抚","加强安保","全城宵禁","电力抢修"]
        resource_keys = ["警力","医疗","物资","电力","通信","交通","供水","燃气","工程抢修","财政资金","网信/舆情","安置保障"]
        cost_data = {k: [] for k in resource_keys}
        for inv in all_interventions:
            prof = intervention_profile(inv)
            for rk in resource_keys:
                cost_data[rk].append(prof["cost"].get(rk, 0))
        fig_cost = go.Figure()
        colors_r = ["#58a6ff","#2ea043","#f85149","#f0883e","#a371f7","#8b949e","#0ea5e9","#22c55e","#f59e0b","#ef4444","#a78bfa","#14b8a6"]
        for i, rk in enumerate(resource_keys):
            fig_cost.add_trace(go.Bar(name=rk, x=all_interventions, y=cost_data[rk], marker_color=colors_r[i]))
        fig_cost.update_layout(barmode="stack", height=220, margin=dict(t=30,b=30), xaxis_tickangle=-25, legend=dict(orientation="h", y=1.1))
        st.plotly_chart(fig_cost, use_container_width=True)
        eff_data = {"calm":[], "order":[], "restore":[]}
        for inv in all_interventions:
            prof = intervention_profile(inv)
            eff_data["calm"].append(round(prof["effect"]["calm"]*100))
            eff_data["order"].append(round(prof["effect"]["order"]*100))
            eff_data["restore"].append(round(prof["effect"]["restore"]*100))
        fig_eff = go.Figure()
        fig_eff.add_trace(go.Bar(name="降焦虑", x=all_interventions, y=eff_data["calm"], marker_color="#f8c447"))
        fig_eff.add_trace(go.Bar(name="压动荡", x=all_interventions, y=eff_data["order"], marker_color="#e34c26"))
        fig_eff.add_trace(go.Bar(name="恢复力", x=all_interventions, y=eff_data["restore"], marker_color="#58a6ff"))
        fig_eff.update_layout(barmode="group", height=220, margin=dict(t=30,b=30), xaxis_tickangle=-25, yaxis_title="效果系数×100", legend=dict(orientation="h", y=1.1))
        st.plotly_chart(fig_eff, use_container_width=True)
    st.caption(f"当前模式：{sim_mode}；随机种子：{seed}（相同输入 + 相同种子可复现演化轨迹）")

reset = st.button("🔁 重置仿真", use_container_width=True)

def request_reset():
    st.session_state.reset_requested = True

if reset:
    request_reset()
if st.session_state.get("reset_requested"):
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    st.rerun()

# ---------------------- 仿真运行 ----------------------
_dbg("16-仿真分支判断")
auto_sig = json.dumps(
    {
        "event": (event or "")[:2000],
        "seed_text_len": len(st.session_state.get("uploaded_doc_text") or ""),
        "sim_mode": sim_mode,
        "seed": int(seed),
        "steps": int(steps),
        "intervention": intervention,
        "enable_swan": bool(enable_swan),
        "hf_params": st.session_state.get("hf_params", {}),
    },
    ensure_ascii=False,
    sort_keys=True,
)[:12000]
should_autorun = bool(client and event and (st.session_state.get("_auto_run_sig") != auto_sig))
if should_autorun:
    try:
        random.seed(int(seed))
        # 智能化判断：是否与政府/社会危机相关
        if force_crisis_mode:
            relevance = {"is_crisis_simulation": True, "reason": "用户强制启用危机推演模式"}
        else:
            relevance = classify_event_relevance(event)
        st.session_state.event_relevance = relevance

        if not relevance["is_crisis_simulation"]:
            # 创意模式：不启动政府/社会资源相关功能，直接完成用户请求
            with st.status("✅ 智能识别为创意模式...", expanded=True) as s:
                s.update(label="🧠 判定与政府社会资源无关，跳过危机推演")
                st.session_state.run_mode = "creative"
                st.session_state.creative_output = gen_creative_response(event)
                s.update(label="✅ 完成", state="complete")
            st.success("已完成！系统判断您的请求与政府/社会危机无关，已直接按您的要求生成内容。")
        else:
            # 危机推演模式：完整流程
            st.session_state.run_mode = "crisis"
            # 根据事件生成动态展示名称（资源与基础设施），体现真人分析般的灵活度
            st.session_state.resource_display_map = select_resource_display_names(event, int(seed))
            st.session_state.infra_display_map = select_infra_display_names(event, int(seed))
            effective_steps = int(steps)
            if sim_mode.startswith("高保真"):
                hp = st.session_state.get("hf_params", {}) or {}
                if hp.get("auto_steps_by_duration") and int(hp.get("total_duration_hours") or 0) > 0:
                    sh = float(hp.get("step_hours") or 12.0)
                    effective_steps = max(1, int(math.ceil(int(hp["total_duration_hours"]) / max(1.0, sh))))
                    if effective_steps > 24:
                        effective_steps = 24
            with st.status("✅ 自动推演中...", expanded=True) as s:
                s.update(label="🌐 检索实证案例")
                doc_text = st.session_state.get("uploaded_doc_text") or ""
                if doc_text and not doc_text.startswith("[") and len(doc_text) > 200:
                    s.update(label="📄 合并用户输入文本到实证上下文")
                    facts = f"""【用户上传的种子材料】
---
{doc_text[:28000]}
---
【以上为种子材料】"""
                    evidence_items = []
                else:
                    facts, evidence_items = get_evidence_items(event)
                st.session_state.event = event
                st.session_state.facts = facts
                st.session_state.evidence_items = evidence_items
                factors, factor_map = extract_evidence_factors_from_items(evidence_items, facts)
                st.session_state.evidence_factor_map = factor_map
                st.session_state.init_factors = factors
                st.session_state.decay_schedule = compute_decay_schedule(
                    evidence_items,
                    int(effective_steps),
                    float(st.session_state.hf_params.get("step_hours", 12.0)),
                    (st.session_state.hf_params.get("decay_scales") or {}),
                )
                st.session_state.time_axis = build_time_axis(
                    int(effective_steps),
                    float(st.session_state.hf_params.get("step_hours", 12.0)),
                    st.session_state.hf_params,
                )

                s.update(label="📊 初始化社会矩阵")
                matrix = init_state(event, facts)
                st.session_state.matrix_history = [matrix]

                s.update(label="🚚 初始化资源调度系统")
                resources = init_resources()
                st.session_state.resources = [resources]
                st.session_state.audit_log = []
                infra = init_infrastructure(factors)
                st.session_state.infra_history = [infra]

                s.update(label="🔗 生成因果链")
                chain = gen_causal(event)
                st.session_state.causal_chain = chain

                if enable_swan:
                    s.update(label="🦢 生成黑天鹅")
                    swan = gen_black_swan(event)
                    st.session_state.swan = swan
                else:
                    st.session_state.swan = "无黑天鹅"

                # 直接全流程演化（去掉逐步推演控制台）
                s.update(label="⏳ 自动演化推演")
                st.session_state.timeline = []
                st.session_state.audit_log = st.session_state.get("audit_log") or []
                for i in range(1, effective_steps + 1):
                    if sim_mode.startswith("高保真"):
                        prev = matrix
                        matrix, resources, infra, audit = mechanistic_step(
                            matrix, resources, factors, intervention, i, infra,
                            st.session_state.hf_params, st.session_state.evidence_items,
                            st.session_state.evidence_factor_map, st.session_state.decay_schedule,
                            st.session_state.time_axis,
                        )
                        st.session_state.audit_log.append(audit)
                        evo = narrate_from_mechanism(event, facts, i, intervention, matrix, prev, resources, audit)
                        st.session_state.timeline.append({"step": i, "data": evo, "audit": audit, "state": matrix.copy(), "resources": resources.copy(), "infra": infra.copy()})
                        st.session_state.infra_history.append(infra.copy())
                    else:
                        evo = evolve_step(event, facts, matrix, i, intervention, resources)
                        st.session_state.timeline.append({"step": i, "data": evo})
                        matrix = update_state(matrix, evo, intervention)
                        resources = update_resources(resources)
                    st.session_state.matrix_history.append(matrix.copy())
                    st.session_state.resources.append(resources.copy())

                s.update(label="📝 生成研判报告")
                if sim_mode.startswith("高保真"):
                    audit_summary = json.dumps(st.session_state.audit_log, ensure_ascii=False)
                    report = gen_report(event, facts + "\n\n[机制审计日志摘要]\n" + audit_summary[:2500], st.session_state.timeline, st.session_state.get("swan",""), matrix, resources)
                else:
                    report = gen_report(event, facts, st.session_state.timeline, st.session_state.get("swan",""), matrix, resources)
                st.session_state.report = report
                _save_to_history("", event, report, st.session_state.matrix_history)
                st.session_state.sim_phase = "complete"
                s.update(label="✅ 推演完成", state="complete")

            st.session_state["_auto_run_sig"] = auto_sig
            st.success("推演完成！已自动完成全流程演化与研判报告生成。")
    except Exception as e:
        st.error(f"仿真流程异常：{str(e)}。请检查 API Key、Base URL 与网络后重试。")
        if "matrix_history" in st.session_state and st.session_state.matrix_history:
            pass
        else:
            st.session_state.timeline = []

# ---------------------- 结果展示 ----------------------
_dbg("17-结果展示分支")
# 创意模式：与政府/社会资源无关，仅展示生成内容
if st.session_state.get("run_mode") == "creative":
    st.divider()
    rel = st.session_state.get("event_relevance") or {}
    st.info(f"🧠 **智能判断**：{rel.get('reason', '与政府社会资源无关')}\n\n以下政府/社会相关板块已跳过：资源调度、基础设施、社会矩阵、干预策略、演化推演等。")
    st.subheader("📝 生成内容")
    st.markdown(f"<div class='report-card'>{st.session_state.get('creative_output', '')}</div>", unsafe_allow_html=True)
    if st.button("🔁 重置", key="reset_creative"):
        request_reset()
        st.rerun()

# 已改为“输入事件后自动全流程演化”，逐步推演控制台已移除

# 危机推演模式：完整展示
elif st.session_state.timeline:
    st.divider()
    st.subheader("📺 实时态势总面板")

    # 1. 状态指标
    m = st.session_state.matrix_history[-1]
    m_cols = st.columns(4)
    labels = ["行政效能","焦虑指数","资源缺口","动荡风险"]
    colors = ["#58a6ff","#f8c447","#f85149","#e34c26"]
    for i,k in enumerate(labels):
        m_cols[i].markdown(f"""
<div class='metric-card'>
{labels[i]}<br>
<h2 style='color:{colors[i]}'>{m[k]}%</h2>
</div>
""", unsafe_allow_html=True)

    # 2. 资源调度面板（展示名称根据事件动态生成，体现灵活分析）
    st.subheader("🚚 资源与兵力调度面板")
    r = st.session_state.resources[-1]
    res_display = st.session_state.get("resource_display_map") or {}
    r_keys = list(r.keys())
    # 按行分布资源卡片，避免资源种类增加后列数不够导致 IndexError
    num_cols = 6
    r_cols = st.columns(num_cols)
    for idx, k in enumerate(r_keys):
        col = r_cols[idx % num_cols]
        label = res_display.get(k, k)
        title = f"{k}（{label}）" if label != k else k
        col.markdown(f"""
<div class='resource-card'>
{title}<br>
<h3>{r[k]}%</h3>
</div>
""", unsafe_allow_html=True)

    # 2.1 基础设施子系统（高保真模式，展示名称根据事件动态生成）
    if sim_mode.startswith("高保真") and st.session_state.get("infra_history"):
        st.subheader("🏗️ 基础设施子系统面板（高保真）")
        inf = st.session_state.infra_history[-1]
        inf_cols = st.columns(4)
        inf_keys = ["电力", "通信", "交通", "供水"]
        inf_display = st.session_state.get("infra_display_map") or {}
        for i, k in enumerate(inf_keys):
            label = inf_display.get(k, k)
            inf_cols[i].markdown(
                f"""
<div class='resource-card'>
{label}<br>
<h3>{inf.get(k, 0)}%</h3>
</div>
""",
                unsafe_allow_html=True,
            )

    # 3. 时序曲线 + 雷达图
    st.subheader("📈 指标演化曲线 & 风险雷达图")
    g1,g2 = st.columns(2)
    with g1:
        df = pd.DataFrame(st.session_state.matrix_history)
        fig = go.Figure()
        for idx,col in enumerate(df.columns):
            fig.add_trace(go.Scatter(y=df[col], name=col, line=dict(color=colors[idx])))
        fig.update_layout(height=350, paper_bgcolor="#161b22", font=dict(color="#e0e0e0"))
        st.plotly_chart(fig, use_container_width=True)
    with g2:
        fig = go.Figure(go.Scatterpolar(r=list(m.values()), theta=list(m.keys()), fill="toself", fillcolor="rgba(31,111,235,0.3)"))
        fig.update_layout(polar=dict(radialaxis=dict(range=[0,100])), height=350, paper_bgcolor="#0e1117", font=dict(color="white"))
        st.plotly_chart(fig, use_container_width=True)

    # 3.1 关键节点摘要
    st.subheader("📌 关键节点摘要")
    df_m = pd.DataFrame(st.session_state.matrix_history)
    step_labels = ["初始"] + [f"第{i}步" for i in range(1, len(df_m))]
    df_m.insert(0, "阶段", step_labels)
    max_risk_step = df_m["动荡风险"].idxmax()
    max_anxiety_step = df_m["焦虑指数"].idxmax()
    key_cols = st.columns(4)
    key_cols[0].metric("动荡风险峰值阶段", step_labels[max_risk_step], f"值 {df_m.loc[max_risk_step, '动荡风险']:.0f}")
    key_cols[1].metric("焦虑指数峰值阶段", step_labels[max_anxiety_step], f"值 {df_m.loc[max_anxiety_step, '焦虑指数']:.0f}")
    key_cols[2].metric("最终行政效能", f"{m['行政效能']}%", "当前步")
    key_cols[3].metric("最终资源缺口", f"{m['资源缺口']}%", "当前步")

    # 3.2 分支对比（如有分支）
    branches = st.session_state.get("scenario_branches") or []
    if branches:
        st.subheader("🌿 分支对比")
        tab_names = ["主路径"] + [b.get("name", f"分支{i+1}") for i, b in enumerate(branches)]
        tabs = st.tabs(tab_names)
        with tabs[0]:
            df_main = pd.DataFrame(st.session_state.matrix_history)
            fig = go.Figure()
            for idx, col in enumerate(df_main.columns):
                fig.add_trace(go.Scatter(y=df_main[col], name=col, line=dict(color=colors[idx % len(colors)])))
            fig.update_layout(height=280, title="主路径指标演化")
            st.plotly_chart(fig, use_container_width=True)
        for i, b in enumerate(branches):
            with tabs[i + 1]:
                df_b = pd.DataFrame(b.get("matrix_history", []))
                if not df_b.empty:
                    fig_b = go.Figure()
                    for idx, col in enumerate(df_b.columns):
                        fig_b.add_trace(go.Scatter(y=df_b[col], name=col, line=dict(color=colors[idx % len(colors)])))
                    fig_b.update_layout(height=280, title=f"分支：{b.get('intervention','')}")
                    st.plotly_chart(fig_b, use_container_width=True)
                st.caption(f"从第 {b.get('fork_step',0)+1} 步起采用「{b.get('intervention','')}」")
                if st.button(f"删除此分支", key=f"del_branch_{i}"):
                    branches.pop(i)
                    st.session_state.scenario_branches = branches
                    st.rerun()

    # 4. 多智能体对抗视图
    st.divider()
    st.subheader("🏛️ 多智能体对抗视图")
    last_evo = st.session_state.timeline[-1]["data"]
    a_cols = st.columns(3)
    with a_cols[0]:
        st.markdown(f"<div class='agent-card'><h4>🏛️ 官方</h4><p>{last_evo.get('official', '—')}</p></div>", unsafe_allow_html=True)
    with a_cols[1]:
        st.markdown(f"<div class='agent-card'><h4>👥 民众</h4><p>{last_evo.get('citizen', '—')}</p></div>", unsafe_allow_html=True)
    with a_cols[2]:
        st.markdown(f"<div class='agent-card'><h4>📺 媒体</h4><p>{last_evo.get('media', '—')}</p></div>", unsafe_allow_html=True)

    # 4.1 与视角对话（支持按步选择）
    st.divider()
    st.subheader("💬 与视角对话")
    st.caption("选择步骤与视角，向该步骤下的该视角提问；系统基于该步的叙事与立场作答")
    if st.session_state.get("perspective_chat_history"):
        if st.button("清空对话记录", key="clear_chat"):
            st.session_state.perspective_chat_history = []
            st.rerun()
    tl = st.session_state.get("timeline") or []
    max_step_idx = len(tl) - 1
    chat_col1, chat_col2, chat_col3 = st.columns([1, 1, 2])
    with chat_col1:
        chat_step_options = [f"第{i+1}步" for i in range(max_step_idx + 1)]
        chat_step_idx = st.selectbox("选择步骤", range(max_step_idx + 1), format_func=lambda x: chat_step_options[x], key="chat_step_select")
    with chat_col2:
        chat_perspective = st.selectbox("选择视角", ["官方", "民众", "媒体", "审计"], key="chat_perspective")
    with chat_col3:
        chat_q = st.chat_input("输入您的问题，如：为什么采取这一策略？民众反应如何？")
    if chat_q:
        nav_map = {"官方": "official", "民众": "citizen", "媒体": "media", "审计": "audit"}
        nav_key = nav_map.get(chat_perspective, "official")
        step_data = tl[chat_step_idx]["data"] if chat_step_idx < len(tl) else last_evo
        narrative = step_data.get(nav_key, "")
        step_state = tl[chat_step_idx].get("state", m) if chat_step_idx < len(tl) and "state" in tl[chat_step_idx] else m
        step_res = tl[chat_step_idx].get("resources", st.session_state.resources[-1]) if chat_step_idx < len(tl) and "resources" in tl[chat_step_idx] else (st.session_state.resources[chat_step_idx] if chat_step_idx < len(st.session_state.resources) else st.session_state.resources[-1])
        ctx = f"事件：{event}\n第{chat_step_idx+1}步状态：{step_state}\n资源：{step_res}"
        reply = chat_with_perspective(chat_perspective, narrative, chat_q, ctx)
        hist = st.session_state.get("perspective_chat_history") or []
        hist.append({"role": "user", "content": chat_q, "perspective": chat_perspective, "step": chat_step_idx + 1})
        hist.append({"role": "assistant", "content": reply, "perspective": chat_perspective, "step": chat_step_idx + 1})
        st.session_state.perspective_chat_history = hist[-20:]
        st.rerun()
    for h in (st.session_state.get("perspective_chat_history") or []):
        with st.chat_message("user" if h["role"] == "user" else "assistant"):
            st.caption(f"【第{h.get('step','')}步·{h.get('perspective','')}视角】")
            st.write(h["content"])

    # 5. 推演回放
    st.divider()
    st.subheader("⏪ 演化推演回放")
    max_step = len(st.session_state.timeline) - 1
    # 自动播放：每轮 rerun 推进一帧，直到结束
    if st.session_state.get("autoplay_running"):
        idx = min(st.session_state.playback_index + 1, max_step)
        st.session_state.playback_index = idx
        if idx >= max_step:
            st.session_state.autoplay_running = False
        time.sleep(1.2)
        st.rerun()
    pb_cols = st.columns([4,1])
    with pb_cols[0]:
        playback = st.slider("回放阶段", 0, max_step, st.session_state.playback_index, key="playback_slider")
    with pb_cols[1]:
        st.markdown("<br>", unsafe_allow_html=True)
        auto = st.button("▶️ 自动播放")
    if auto:
        st.session_state.autoplay_running = True
        st.session_state.playback_index = 0
        st.rerun()
    st.session_state.playback_index = playback
    play_data = st.session_state.timeline[playback]["data"]
    # 场景分支：从当前步创建分支
    with st.expander("🌿 从本步创建分支（对比不同干预路径）", expanded=False):
        branch_interv = st.selectbox("假设从本步起改用干预策略", ["无","紧急物资投放","媒体安抚","加强安保","全城宵禁","电力抢修"], key="branch_intervention")
        branch_name = st.text_input("分支名称（可选）", value=f"第{playback+1}步起→{branch_interv}", key="branch_name")
        if st.button("创建并推演分支", key="create_branch_btn"):
            with st.spinner("正在推演分支..."):
                branch_data = run_branch_from_step(playback, branch_interv, event, st.session_state.get("facts",""))
                branch_data["name"] = branch_name or f"分支-第{playback+1}步起→{branch_interv}"
                st.session_state.scenario_branches = st.session_state.get("scenario_branches") or []
                st.session_state.scenario_branches.append(branch_data)
                st.success(f"分支「{branch_name}」推演完成，可在下方查看对比")
                st.rerun()
    st.info(f"📌 第 {playback+1} 阶段 演化内容")
    st.markdown(f"**官方**：{play_data.get('official', '—')}")
    st.markdown(f"**民众**：{play_data.get('citizen', '—')}")
    st.markdown(f"**媒体**：{play_data.get('media', '—')}")
    st.markdown(f"**审计**：{play_data.get('audit', '—')}")

    # 5.1 实体与知识图
    with st.expander("🕸️ 实体与知识图谱", expanded=False):
        if st.button("提取实体与关系", key="extract_entities_btn"):
            with st.spinner("正在抽取..."):
                eg = extract_entity_graph(event, st.session_state.get("facts",""))
                st.session_state.entity_graph = eg
        eg = st.session_state.get("entity_graph") or {}
        if eg.get("entities") or eg.get("relations"):
            e_col1, e_col2 = st.columns(2)
            with e_col1:
                st.markdown("**实体**")
                for ent in eg.get("entities", []):
                    st.write(f"- {ent.get('name','')}（{ent.get('type','')}）")
            with e_col2:
                st.markdown("**关系**")
                for rel in eg.get("relations", []):
                    st.write(f"- {rel.get('source','')} → {rel.get('target','')}：{rel.get('relation','')}")
            # 简单网络图（需 networkx，可选）
            try:
                import networkx as nx
                G = nx.DiGraph()
                for e in eg.get("entities", []):
                    G.add_node(e.get("name", ""), type=e.get("type", ""))
                for r in eg.get("relations", []):
                    G.add_edge(r.get("source",""), r.get("target",""), label=r.get("relation",""))
                if G.number_of_nodes() > 0:
                    pos = nx.spring_layout(G, seed=42)
                    edge_trace = go.Scatter(x=[], y=[], line=dict(width=1, color="#888"), hoverinfo="text", mode="lines")
                    node_trace = go.Scatter(x=[], y=[], text=[], mode="markers+text", marker=dict(size=20, color="#58a6ff"), textposition="top center")
                    for e in G.edges():
                        x0, y0 = pos[e[0]]
                        x1, y1 = pos[e[1]]
                        edge_trace["x"] += (x0, x1, None)
                        edge_trace["y"] += (y0, y1, None)
                    for n in G.nodes():
                        node_trace["x"] += (pos[n][0],)
                        node_trace["y"] += (pos[n][1],)
                        node_trace["text"] += (n,)
                    fig_g = go.Figure(data=[edge_trace, node_trace])
                    fig_g.update_layout(title="知识图谱", showlegend=False, height=300)
                    st.plotly_chart(fig_g, use_container_width=True)
            except ImportError:
                st.caption("安装 networkx 可显示图谱：pip install networkx")

    # 6. 因果链 + 黑天鹅
    st.divider()
    st.subheader("🔗 深度因果演化链")
    st.markdown(f"<div class='logic-box'>{st.session_state.causal_chain}</div>", unsafe_allow_html=True)
    st.error(f"🦢 黑天鹅事件：{st.session_state.swan}")
    st.warning(f"🛡️ 系统审计：{last_evo.get('audit', '—')}")
    # 以下审计/映射/衰减/时间轴信息仅后台保留，不在界面展示（按需求隐藏）
    if sim_mode.startswith("高保真") and st.session_state.get("evidence_items"):
        with st.expander("📎 证据清单（可引用）", expanded=False):
            st.json(st.session_state.evidence_items)
    # （隐藏）证据→因子映射 / 冲击衰减曲线 / 时间轴 的界面展示
    if sim_mode.startswith("高保真") and st.session_state.get("hf_params"):
        hp = st.session_state.get("hf_params") or {}
        if hp.get("auto_steps_by_duration") and int(hp.get("total_duration_hours") or 0) > 0:
            st.caption(f"时间设定（高保真）：总时长={int(hp['total_duration_hours'])}h，每步={float(hp.get('step_hours', 12.0))}h（步数自动计算并可能上限为 24）")

    # 7. 报告 + 导出
    st.divider()
    st.subheader("📝 国家级战略研判报告")
    st.markdown(f"<div class='report-card'>{st.session_state.report}</div>", unsafe_allow_html=True)

    # 若存在上传文档的全文内容分析报告，则一并展示，体现“基于全部内容的分析研判”
    if st.session_state.get("uploaded_doc_report"):
        st.subheader("📄 文档内容分析研判（基于上传全文）")
        st.markdown(f"<div class='report-card'>{st.session_state.uploaded_doc_report}</div>", unsafe_allow_html=True)

    # 导出区
    d_cols = st.columns(4)
    with d_cols[0]:
        doc_text = st.session_state.get("uploaded_doc_text") or ""
        doc_summary = doc_text[:1500] + ("..." if len(doc_text) > 1500 else "") if doc_text and not doc_text.startswith("[") else ""
        json_data = json.dumps({
            "event": event, "facts": st.session_state.facts,
            "matrix": st.session_state.matrix_history,
            "timeline": st.session_state.timeline,
            "report": st.session_state.report,
            "seed_document_summary": doc_summary,
            "audit_log": st.session_state.get("audit_log", []),
            "evidence_items": st.session_state.get("evidence_items", []),
            "evidence_factor_map": st.session_state.get("evidence_factor_map", {}),
            "decay_schedule": st.session_state.get("decay_schedule", {}),
            "time_axis": st.session_state.get("time_axis", {}),
            "infra_history": st.session_state.get("infra_history", []),
            "hf_params": st.session_state.get("hf_params", {}),
            "resource_display_map": st.session_state.get("resource_display_map", {}),
            "infra_display_map": st.session_state.get("infra_display_map", {}),
            "uploaded_doc_analysis": st.session_state.get("uploaded_doc_analysis", {}),
            "uploaded_doc_report": st.session_state.get("uploaded_doc_report", ""),
        }, ensure_ascii=False, indent=2)
        fname_json = "推演数据包.json"
        st.download_button("💾 导出JSON数据包", json_data, fname_json, use_container_width=True)
    with d_cols[1]:
        csv_df = pd.DataFrame(st.session_state.matrix_history)
        csv_df.insert(0, "阶段", ["初始"] + [f"第{i}步" for i in range(1, len(csv_df))])
        fname_csv = "指标演化.csv"
        st.download_button("📊 导出指标CSV", csv_df.to_csv(index=False).encode("utf-8-sig"), fname_csv, "text/csv", use_container_width=True)
    with d_cols[2]:
        pdf_report = st.session_state.report
        pdf_doc_appendix = (st.session_state.get("uploaded_doc_text") or "")[:4000]
        if pdf_doc_appendix and pdf_doc_appendix.startswith("["):
            pdf_doc_appendix = ""
        st.download_button("📄 导出PDF报告", export_pdf(pdf_report, pdf_doc_appendix), "研判报告.pdf", use_container_width=True)
    with d_cols[3]:
        if st.button("🔁 重置仿真", use_container_width=True, key="reset_bottom"):
            request_reset()
            st.rerun()

elif not client:
    st.warning("请在左侧填写 API Key 后启动仿真")
_dbg("18-脚本执行完毕")
